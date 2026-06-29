"""Dynamic API Agent — Nango-free.

A self-contained agent that lets a user connect to ANY external tool/API
through natural language. No hardcoded provider list, no Nango. The local
Ollama LLM does the reasoning (tool identification, doc extraction, action
planning), this module does the I/O (web search, HTTP, encrypted storage).

Pipeline for a single turn:

    user prompt
        │
        ├─► identify_tool          ── LLM picks the tool from pure reasoning
        │
        ├─► lookup_or_fetch_docs   ── DB cache → web search → LLM extract
        │
        ├─► load_connection        ── existing creds? OAuth refresh needed?
        │        │
        │        └─► if missing → ask_user_creds  (return needs_credentials)
        │
        ├─► plan_action            ── LLM converts prompt to {method,path,…}
        │
        └─► execute_action         ── raw HTTP with injected credentials

Each step emits a structured Thought / Action / Action_Input / Summary so
the frontend can render the agent's reasoning, not just the final answer.
"""

from __future__ import annotations

import json
import re
import time
from datetime import datetime, timedelta
from typing import Any, Callable, Dict, List, Optional, Tuple

import requests
from sqlalchemy.orm import Session
from sqlalchemy.orm.exc import StaleDataError

from app.core.config import settings
from app.core.logger import logger
from app.core.security import decrypt_api_key, encrypt_api_key
from app.db.models import (
    DynamicAgentRunLog,
    DynamicToolConnection,
    ToolDefinition,
    UserAPIKey,
)
from app.services.llm_provider import LLMProvider

# ---------------------------------------------------------------- LLM prompts

_TOOL_IDENTIFY_SYSTEM = __import__("base64").b64decode("WW91IGFyZSB0aGUgdG9vbC1yb3V0ZXIgZm9yIGEgZHluYW1pYyBBUEkgYWdlbnQuCgpUaGUgdXNlciBqdXN0IHNhaWQgc29tZXRoaW5nLiBEZWNpZGUgd2hpY2ggRVhURVJOQUwgQVBJL1RPT0wgdGhleSB3YW50IHRvCnVzZS4gVGhlcmUgaXMgTk8gY2F0YWxvZyDigJQgeW91IG11c3QgcmVhc29uIGZyb20gdGhlIHByb21wdCdzIGludGVudC4KCkNvbW1vbiB0b29scyB5b3UnbGwgc2VlOiBnaXRodWIsIG5vdGlvbiwgZ21haWwsIHNsYWNrLCBnb29nbGUtY2FsZW5kYXIsCmdvb2dsZS1kcml2ZSwgZ29vZ2xlLXNoZWV0cywgbGluZWFyLCBhc2FuYSwgamlyYSwgdHJlbGxvLCBzdHJpcGUsCnJhem9ycGF5LCBwYXlwYWwsIG9wZW5haSwgYW50aHJvcGljLCBodWJzcG90LCBzYWxlc2ZvcmNlLCBzaG9waWZ5LAptYWlsY2hpbXAsIHNlbmRncmlkLCB0d2lsaW8sIGRpc2NvcmQsIHRlbGVncmFtLCB6b29tLCBmaWdtYSwgYWlydGFibGUsCmF3cywgZ2NwLCBhenVyZS4KQnV0IEFOWSBwdWJsaWMgUkVTVCB0b29sIGlzIHZhbGlkIOKAlCBwaWNrIHRoZSBjYW5vbmljYWwgbG93ZXJjYXNlIG5hbWUKdGhlIHdvcmxkIHVzZXMgKGUuZy4gImdpdGh1YiIgbm90ICJHaXRIdWIgSW5jLiIsICJnb29nbGUtY2FsZW5kYXIiIG5vdAoiZ2NhbCIpLgoKUnVsZXM6Ci0gSWYgdGhlIHVzZXIgbmFtZWQgYSB0b29sIGRpcmVjdGx5ICgiY29ubmVjdCBnaXRodWIiKSDihpIgdGhhdCdzIHRoZSB0b29sLgotIElmIHRoZXkgZGVzY3JpYmVkIGFuIGFjdGlvbiAoInNlbmQgYW4gZW1haWwiKSDihpIgcGljayB0aGUgbW9zdCBjb21tb24gdG9vbAogIGZvciB0aGF0IGFjdGlvbiAoZ21haWwgZm9yICJzZW5kIGVtYWlsIiwgc2xhY2sgZm9yICJwb3N0IGEgbWVzc2FnZSIgd2l0aAogIG5vIG90aGVyIGNvbnRleHQsIGV0Yy4pLgotICJ0b29sIiBpcyB0aGUgY2Fub25pY2FsIGxvd2VyY2FzZSBpZGVudGlmaWVyLCBoeXBoZW4tc2VwYXJhdGVkIGZvcgogIG11bHRpLXdvcmQgbmFtZXMuCi0gSU1QT1JUQU5UIOKAlCBjbG91ZCBwcm92aWRlcnMgc3RheSBhcyBPTkUgdG9vbDogZm9yIEFOWVRISU5HIG9uIEFtYXpvbgogIFdlYiBTZXJ2aWNlcyAoRUMyLCBTMywgUkRTLCBMYW1iZGEsIElBTSwg4oCmKSB0aGUgdG9vbCBpcyBleGFjdGx5ICJhd3MiLgogIE5ldmVyIHNwbGl0IGludG8gImF3cy1lYzIiLCAiYXdzLXMzIiwgImFtYXpvbi13ZWItc2VydmljZXMiLCAiYXdzLWNsaSIsCiAgZXRjLiBTYW1lIHJ1bGUgZm9yICJnY3AiIGFuZCAiYXp1cmUiLgotICJpbnRlbnQiIGlzIHdoYXQgdGhlIHVzZXIgd2FudHMgdG8gRE86ICJjb25uZWN0IiAoc3RhcnQgYXV0aCkgLwogICJhY3Rpb24iIChleGVjdXRlIHNvbWV0aGluZyB0aGV5IGFscmVhZHkgYXV0aGVkKSAvICJhbWJpZ3VvdXMiLgotIElmIHlvdSBnZW51aW5lbHkgY2FuJ3QgcGljayBhIHRvb2wgKGUuZy4gImhlbGxvIiksIHJldHVybiB0b29sPW51bGwuCgotICJ3YW50c19kb2NfaW1wb3J0IjogdHJ1ZSBPTkxZIHdoZW4gdGhlIHVzZXIgaXMgYXNraW5nIHRvIFNFVCBVUCAvIE9OQk9BUkQgLwogIFRFQUNIIC8gQUREIGEgdG9vbCBieSBnaXZpbmcgeW91IGl0cyBET0NVTUVOVEFUSU9OIG9yIEFQSSBTUEVDICh1c3VhbGx5IGEKICBsaW5rIG9yIGEgcGhyYXNlIGxpa2UgInNldCB1cCBYIGZyb20gdGhlc2UgZG9jcyIsICJ1c2UgdGhpcyBhcGkgc3BlYyIsCiAgImhlcmUgYXJlIHRoZSBkb2NzIGZvciBYIiwgImltcG9ydCBYIGZyb20gPGxpbms+IikuIEl0IG1lYW5zOiAicmVhZCB0aGlzCiAgZG9jIGFuZCBsZWFybiB0aGUgdG9vbCBmcm9tIGl0LiIKICBTZXQgaXQgRkFMU0UgZm9yIG9yZGluYXJ5IGFjdGlvbiByZXF1ZXN0cywgRVZFTiBJRiB0aGUgcHJvbXB0IGNvbnRhaW5zIGEKICBVUkwg4oCUIGEgbGluayB0aGF0IGlzIGp1c3QgQ09OVEVOVCBvZiB0aGUgdGFzayBpcyBub3QgYSBkb2MuCiAgRXhhbXBsZXM6CiAgICDigKIgInNldCB1cCBhY21lIGZyb20gaHR0cHM6Ly9hY21lLmRldi9vcGVuYXBpLmpzb24iICAgICAg4oaSIHRydWUKICAgIOKAoiAiaGVyZSBhcmUgdGhlIGxpbmtlZGluIGFwaSBkb2NzOiA8bGluaz4sIGFkZCBsaW5rZWRpbiIg4oaSIHRydWUKICAgIOKAoiAicG9zdCBvbiBsaW5rZWRpbiBzaGFyaW5nIGh0dHBzOi8vbXlzaXRlLmNvbS9ibG9nIiAgICAg4oaSIGZhbHNlICh0aGUgVVJMCiAgICAgIGlzIGNvbnRlbnQgdG8gcG9zdCwgbm90IGEgZG9jIHRvIGxlYXJuIGZyb20pCiAgICDigKIgImNyZWF0ZSBhIGdpdGh1YiBpc3N1ZSBsaW5raW5nIHRvIGh0dHBzOi8veC5jb20vYnVnIiAgIOKGkiBmYWxzZQoKUmVzcG9uZCB3aXRoIFNUUklDVCBKU09OIG9ubHksIG5vIHByb3NlLCBubyBtYXJrZG93biBmZW5jZXM6Cgp7CiAgInRvb2wiOiAiPGxvd2VyY2FzZSBjYW5vbmljYWwgbmFtZT4iIHwgbnVsbCwKICAiaW50ZW50IjogImNvbm5lY3QiIHwgImFjdGlvbiIgfCAiYW1iaWd1b3VzIiwKICAid2FudHNfZG9jX2ltcG9ydCI6IHRydWUgfCBmYWxzZSwKICAiY29uZmlkZW5jZSI6IDwwLi4xPiwKICAicmVhc29uIjogIjxvbmUgc2VudGVuY2U+Igp9Cg==").decode()

_DOCS_EXTRACT_SYSTEM = __import__("base64").b64decode("WW91IHJlYWQgcmF3IEFQSSBkb2N1bWVudGF0aW9uIHBhZ2VzIGFuZCBleHRyYWN0CnRoZSBzdHJ1Y3R1cmVkIGZpZWxkcyB0aGUgYWdlbnQgbmVlZHMgdG8gY2FsbCB0aGlzIEFQSS4KCk91dHB1dCBTVFJJQ1QgSlNPTiBvbmx5OgoKewogICJiYXNlX3VybCI6ICAgIjxodHRwczovL2FwaS5leGFtcGxlLmNvbT4iLAogICJhdXRoX3R5cGUiOiAgIkFQSV9LRVkiIHwgIkJFQVJFUiIgfCAiT0FVVEgyIiB8ICJPQVVUSDEiIHwgIkJBU0lDIiB8ICJQQVQiLAogICJhdXRoX2NvbmZpZyI6IHsKICAgIC8vIEZpbGwgT05MWSB0aGUgZmllbGRzIHRoYXQgYXBwbHkgdG8gYXV0aF90eXBlOgogICAgImhlYWRlcl9uYW1lIjogICAgICAgICJBdXRob3JpemF0aW9uIiwgICAgICAgLy8gZm9yIEFQSV9LRVkgLyBCRUFSRVIKICAgICJjcmVkZW50aWFsX3ByZWZpeCI6ICAiQmVhcmVyICIsICAgICAgICAgICAgIC8vIGZvciBBUElfS0VZIC8gQkVBUkVSCiAgICAicXVlcnlfcGFyYW0iOiAgICAgICAgImFwaV9rZXkiLCAgICAgICAgICAgICAvLyBhbHQgZm9yIEFQSV9LRVkKICAgICJvYXV0aF9hdXRob3JpemVfdXJsIjoiaHR0cHM6Ly/igKYvYXV0aG9yaXplIiwgLy8gZm9yIE9BVVRIMgogICAgIm9hdXRoX3Rva2VuX3VybCI6ICAgICJodHRwczovL+KApi90b2tlbiIsICAgICAvLyBmb3IgT0FVVEgyCiAgICAiZGVmYXVsdF9zY29wZXMiOiAgICAgInJlcG8scmVhZDp1c2VyIiwgICAgICAvLyBmb3IgT0FVVEgyCiAgICAiY2FsbGJhY2tfdXJsX2hpbnQiOiAgImh0dHBzOi8veW91ci1hcHAuY29tL29hdXRoL2NhbGxiYWNrIiwKICAgICJjcmVkZW50aWFsX2ZpZWxkX292ZXJyaWRlcyI6IHsgICAgICAgICAgICAgIC8vIGZyaWVuZGx5IGxhYmVscyBmb3IgdGhlCiAgICAgICJ1c2VybmFtZSI6IHsibGFiZWwiOiAiS2V5IElEIn0sICAgICAgICAgICAvLyBjcmVkZW50aWFsIGZvcm0gd2hlbiB0aGUKICAgICAgInBhc3N3b3JkIjogeyJsYWJlbCI6ICJLZXkgU2VjcmV0In0gICAgICAgIC8vIGRvY3MgY2FsbCB0aGVtIHNvbWV0aGluZwogICAgfSAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgICAgLy8gc3BlY2lmaWMgKHNlZSBydWxlIGJlbG93KQogIH0sCiAgImVuZHBvaW50cyI6IHsKICAgICI8dmVyYl9uYW1lPiI6IHsKICAgICAgIm1ldGhvZCI6ICJHRVQiIHwgIlBPU1QiIHwgIlBVVCIgfCAiUEFUQ0giIHwgIkRFTEVURSIsCiAgICAgICJwYXRoIjogICAiL3BhdGgvdW5kZXIvYmFzZV91cmwiLAogICAgICAiZGVzY3JpcHRpb24iOiAib25lLWxpbmUgc3VtbWFyeSIsCiAgICAgICJwYXJhbXMiOiBudWxsIHwgeyAiPHBhcmFtPiI6ICI8ZGVzY3JpcHRpb24+IiB9LAogICAgICAiYm9keSI6ICAgbnVsbCB8IHsgIjxmaWVsZD4iOiAiPGRlc2NyaXB0aW9uPiIgfQogICAgfSwKICAgICI8bW9yZV92ZXJicz4iOiB7IOKApiB9CiAgfSwKICAicmF0ZV9saW1pdHMiOiBudWxsIHwgewogICAgInJlcXVlc3RzX3Blcl9taW51dGUiOiA8aW50PiB8IG51bGwsCiAgICAicmVxdWVzdHNfcGVyX2hvdXIiOiAgIDxpbnQ+IHwgbnVsbCwKICAgICJyZXF1ZXN0c19wZXJfZGF5IjogICAgPGludD4gfCBudWxsLAogICAgIm5vdGVzIjogICAgICAgICAgICAgICAiPG9uZSBzaG9ydCBzZW50ZW5jZSDigJQgYnVyc3QgbGltaXRzLCB0aWVycywgZXRjLj4iCiAgfSwKICAiZXhhbXBsZXMiOiBudWxsIHwgWwogICAgewogICAgICAibGFuZ3VhZ2UiOiAgICAiY3VybCIgfCAicHl0aG9uIiB8ICJqYXZhc2NyaXB0IiB8ICJzaGVsbCIgfCAiaHR0cCIsCiAgICAgICJ0aXRsZSI6ICAgICAgICI8c2hvcnQgbGFiZWwgZS5nLiAnRmV0Y2ggdXNlcicgPiIsCiAgICAgICJjb2RlIjogICAgICAgICI8c2luZ2xlIGNvZGUgYmxvY2sg4oCUIGtlZXAgdW5kZXIgMzAgbGluZXM+IgogICAgfQogIF0sCiAgInF1aXJrcyI6IG51bGwgfCBbCiAgICAiPHNob3J0IGltcGVyYXRpdmUgcnVsZSB0aGUgY2FsbGVyIE1VU1QgZm9sbG93LCB0YWtlbiBmcm9tIHRoZSBkb2NzPiIKICBdLAogICJkb2NzX3VybCI6ICI8Y2Fub25pY2FsIGRvY3MgVVJMPiIKfQoKUnVsZXM6Ci0gImJhc2VfdXJsIiBNVVNUIGJlIHRoZSBBUEkgaG9zdG5hbWUgKGUuZy4gaHR0cHM6Ly9hcGkuZ2l0aHViLmNvbSksIE5PVAogIHRoZSBodW1hbiBkb2NzIHBhZ2UuCi0gUGF0aHMgYXJlIFJFTEFUSVZFIHRvIGJhc2VfdXJsIChzdGFydCB3aXRoICIvIikuIE5ldmVyIGluY2x1ZGUgdGhlIGhvc3QKICBpbnNpZGUgYHBhdGhgLgotIEV4dHJhY3QgRVZFUlkgZW5kcG9pbnQgeW91IGNhbiBmaW5kIGluIHRoZSBkb2N1bWVudGF0aW9uIOKAlCBkbyBub3QgbGltaXQKICB5b3Vyc2VsZiB0byBhIGZldyAiY29tbW9uIiBvbmVzLiBJbmNsdWRlIGFsbCBHRVQsIFBPU1QsIFBVVCwgUEFUQ0gsIGFuZAogIERFTEVURSBlbmRwb2ludHMgbWVudGlvbmVkLiBUaGUgZ29hbCBpcyBtYXhpbXVtIGNvdmVyYWdlLgotIEZvciBPQVVUSDIgdG9vbHMsIHBvcHVsYXRlIG9hdXRoX2F1dGhvcml6ZV91cmwgKyBvYXV0aF90b2tlbl91cmwgZXZlbgogIGlmIHRoZSBkb2NzIG9ubHkgbWVudGlvbiB0aGVtIGJyaWVmbHkuCi0gY3JlZGVudGlhbF9maWVsZF9vdmVycmlkZXMg4oCUIEFMV0FZUyBzZXQgdGhpcyBmb3IgQkFTSUMgYW5kIEFQSV9LRVkgLyBCRUFSRVIKICAvIFBBVCBhdXRoLiBSZWFkIHRoZSBkb2NzIGFuZCB1c2UgdGhlIEVYQUNUIG5hbWVzIHRoZSBwcm92aWRlciB1c2VzIGZvciBpdHMKICBjcmVkZW50aWFscywgc28gdGhlIHVzZXIgc2VlcyB0aGUgc2FtZSB3b3JkcyBhcyBpbiB0aGUgZGFzaGJvYXJkIGluc3RlYWQgb2YKICBnZW5lcmljICJVc2VybmFtZS9QYXNzd29yZCIuIFRoZSBvdmVycmlkYWJsZSBmaWVsZCBuYW1lcyBhcmU6ICJ1c2VybmFtZSIgKwogICJwYXNzd29yZCIgKEJBU0lDKSwgYW5kICJzZWNyZXQiIChBUElfS0VZIC8gQkVBUkVSIC8gUEFUKS4gRXhhbXBsZXMgZnJvbSByZWFsCiAgZG9jczogUmF6b3JwYXkgQkFTSUMg4oaSIHVzZXJuYW1lPSJLZXkgSUQiLCBwYXNzd29yZD0iS2V5IFNlY3JldCI7IFR3aWxpbyBCQVNJQwogIOKGkiB1c2VybmFtZT0iQWNjb3VudCBTSUQiLCBwYXNzd29yZD0iQXV0aCBUb2tlbiI7IFN0cmlwZSBBUElfS0VZIOKGkgogIHNlY3JldD0iU2VjcmV0IEtleSIuIE9ubHkgZmFsbCBiYWNrIHRvIG9taXR0aW5nIGl0IGlmIHRoZSBkb2NzIGdlbnVpbmVseSB1c2UKICB0aGUgbGl0ZXJhbCB3b3JkcyB1c2VybmFtZS9wYXNzd29yZC4KLSAicmF0ZV9saW1pdHMiIOKAlCBvbmx5IHBvcHVsYXRlIGlmIHRoZSBkb2NzIE1FTlRJT04gc3BlY2lmaWMgbnVtYmVycy4KICBEb24ndCBpbnZlbnQuIElmIHRoZSBkb2NzIG9ubHkgc2F5ICJyYXRlIGxpbWl0cyBhcHBseSIsIHNldCB0byBudWxsLgotICJleGFtcGxlcyIg4oCUIGluY2x1ZGUgMS0zIHJlYWwgY29kZSBibG9ja3MgeW91IHNhdyBpbiB0aGUgZG9jcy4gRG9uJ3QKICBmYWJyaWNhdGU7IGlmIG5vIGNvZGUgc2FtcGxlcyBhcHBlYXJlZCwgc2V0IHRvIG51bGwuCi0gInF1aXJrcyIg4oCUIGNhcHR1cmUgcHJvdmlkZXItc3BlY2lmaWMgUlVMRVMgdGhhdCBjaGFuZ2UgaG93IGEgY2FsbCBtdXN0IGJlCiAgbWFkZSwgYnV0IE9OTFkgaWYgdGhlIGRvY3Mgc3RhdGUgdGhlbS4gRXhhbXBsZXMgb2Ygd2hhdCB0byBsb29rIGZvcjoKICAgIOKAoiBtb25leS9hbW91bnQgdW5pdHMgKCJhbW91bnQgaXMgaW4gcGFpc2UgLyBjZW50cyDigJQgdGhlIHNtYWxsZXN0IGN1cnJlbmN5CiAgICAgIHVuaXQiLCAic2VuZCBhbW91bnRzIGluIGNlbnRzIikKICAgIOKAoiByZXF1aXJlZCBoZWFkZXJzIChpZGVtcG90ZW5jeSBrZXlzLCBBUEkgdmVyc2lvbiBoZWFkZXJzKQogICAg4oCiIHRpbWVzdGFtcC9kYXRlIGZvcm1hdHMgKHVuaXggc2Vjb25kcyB2cyBJU08tODYwMSkKICAgIOKAoiBwYWdpbmF0aW9uIHN0eWxlLCBtYXggcGFnZSBzaXplLCBtYW5kYXRvcnkgcXVlcnkgcGFyYW1zCiAgICDigKIgaWQgZm9ybWF0cywgcmVnaW9uL2hvc3QgcmVxdWlyZW1lbnRzCiAgRWFjaCBxdWlyayBpcyBPTkUgc2hvcnQgaW1wZXJhdGl2ZSBzZW50ZW5jZS4gRG8gTk9UIGludmVudCBydWxlcyB0aGF0IGFyZW4ndAogIGluIHRoZSBkb2NzOyBpZiB0aGVyZSBhcmUgbm9uZSwgc2V0IHF1aXJrcyB0byBudWxsLiBUaGlzIGlzIGNyaXRpY2FsIGZvcgogIHBheW1lbnQgZ2F0ZXdheXMgd2hlcmUgYW1vdW50cyBhcmUgaW4gdGhlIHNtYWxsZXN0IGN1cnJlbmN5IHVuaXQuCi0gRU5EUE9JTlQgQk9ESUVTIEFSRSBNQU5EQVRPUlk6IGZvciBldmVyeSBQT1NUL1BVVC9QQVRDSCBlbmRwb2ludCwgZmlsbCBpdHMKICBgYm9keWAgd2l0aCBFVkVSWSByZXF1ZXN0IGZpZWxkIHRoZSBkb2NzIGxpc3QgKG5hbWUg4oaSIHNob3J0IGRlc2NyaXB0aW9uKS4KICBBbiBlbXB0eS9udWxsIGJvZHkgb24gYSB3cml0ZSBlbmRwb2ludCBpcyBhbG1vc3QgYWx3YXlzIFdST05HIOKAlCBsb29rIGhhcmRlcgogIGluIHRoZSBkb2NzIGZvciB0aGUgcmVxdWVzdCBwYXJhbWV0ZXJzLiBUaGUgYWdlbnQgcmVqZWN0cyBjYWxscyB3aXRoIGZpZWxkcwogIG5vdCBpbiB0aGlzIGxpc3QsIHNvIGNvbXBsZXRlbmVzcyBoZXJlIGlzIHdoYXQgbWFrZXMgd3JpdGUgYWN0aW9ucyB3b3JrLgoKRVhBTVBMRSDigJQgc3R1ZHkgdGhlIHNoYXBlLiBOb3RlIHRoZSBGVUxMWS1wb3B1bGF0ZWQgYm9keSBhbmQgdGhlIGNyZWRlbnRpYWwKbGFiZWxzICh0aGlzIGlzIGEgQkFTSUMtYXV0aCBwYXltZW50IEFQSSk6Cgp7CiAgImJhc2VfdXJsIjogImh0dHBzOi8vYXBpLmFjbWVwYXkuY29tL3YxIiwKICAiYXV0aF90eXBlIjogIkJBU0lDIiwKICAiYXV0aF9jb25maWciOiB7CiAgICAiY3JlZGVudGlhbF9maWVsZF9vdmVycmlkZXMiOiB7CiAgICAgICJ1c2VybmFtZSI6IHsibGFiZWwiOiAiS2V5IElEIn0sCiAgICAgICJwYXNzd29yZCI6IHsibGFiZWwiOiAiS2V5IFNlY3JldCJ9CiAgICB9CiAgfSwKICAiZW5kcG9pbnRzIjogewogICAgImNyZWF0ZV9wYXltZW50X2xpbmsiOiB7CiAgICAgICJtZXRob2QiOiAiUE9TVCIsCiAgICAgICJwYXRoIjogIi9wYXltZW50X2xpbmtzIiwKICAgICAgImRlc2NyaXB0aW9uIjogIkNyZWF0ZSBhIHBheW1lbnQgbGluayIsCiAgICAgICJwYXJhbXMiOiBudWxsLAogICAgICAiYm9keSI6IHsKICAgICAgICAiYW1vdW50IjogImludGVnZXIsIHNtYWxsZXN0IGN1cnJlbmN5IHVuaXQgKGUuZy4gcGFpc2UpIiwKICAgICAgICAiY3VycmVuY3kiOiAiSVNPIGNvZGUsIGUuZy4gSU5SIiwKICAgICAgICAiZGVzY3JpcHRpb24iOiAidGV4dCBzaG93biB0byB0aGUgY3VzdG9tZXIiLAogICAgICAgICJjdXN0b21lciI6ICJvYmplY3Q6IHtuYW1lLCBlbWFpbCwgY29udGFjdH0iLAogICAgICAgICJub3RpZnkiOiAib2JqZWN0OiB7c21zOiBib29sLCBlbWFpbDogYm9vbH0iLAogICAgICAgICJyZW1pbmRlcl9lbmFibGUiOiAiYm9vbGVhbiIKICAgICAgfQogICAgfQogIH0sCiAgInJhdGVfbGltaXRzIjogbnVsbCwKICAiZXhhbXBsZXMiOiBudWxsLAogICJxdWlya3MiOiBbImFtb3VudCBpcyBpbiB0aGUgc21hbGxlc3QgY3VycmVuY3kgdW5pdCDigJQgbXVsdGlwbHkgcnVwZWVzIGJ5IDEwMCJdLAogICJkb2NzX3VybCI6ICJodHRwczovL2RvY3MuYWNtZXBheS5jb20iCn0K").decode()

_ACTION_PLAN_SYSTEM = __import__("base64").b64decode("WW91IHRyYW5zbGF0ZSBhIG5hdHVyYWwtbGFuZ3VhZ2UgaW5zdHJ1Y3Rpb24gaW50bwpPTkUgSFRUUCBjYWxsIGFnYWluc3QgdGhlIGNvbm5lY3RlZCBwcm92aWRlci4KCllvdSByZWNlaXZlOgogIC0gdG9vbDogICAgICB3aGljaCBwcm92aWRlciBpcyBjb25uZWN0ZWQKICAtIGJhc2VfdXJsOiAgdGhlIHByb3ZpZGVyJ3MgQVBJIGhvc3QKICAtIGVuZHBvaW50czoga25vd24gdmVyYnMgeW91IGNhbiBwaWNrIGZyb20gKGRvIHByZWZlciB0aGVzZSBvdmVyIGludmVudGluZykKICAtIHByb21wdDogICAgd2hhdCB0aGUgdXNlciB3YW50cyB0byBkbwoKUmVzcG9uZCB3aXRoIFNUUklDVCBKU09OIG9ubHk6Cgp7CiAgIm1ldGhvZCI6ICAgIkdFVCIgfCAiUE9TVCIgfCAiUFVUIiB8ICJQQVRDSCIgfCAiREVMRVRFIiwKICAiZW5kcG9pbnQiOiAiL3BhdGgvdW5kZXIvYmFzZV91cmwiLAogICJwYXJhbXMiOiAgIG51bGwgfCB7IOKApiB9LAogICJib2R5IjogICAgIG51bGwgfCB7IOKApiB9LAogICJzdW1tYXJ5IjogICI8b25lIHNob3J0IHNlbnRlbmNlIOKAlCB3aGF0IHRoaXMgY2FsbCB3aWxsIGRvPiIKfQoKUnVsZXM6Ci0gImVuZHBvaW50IiBNVVNUIHN0YXJ0IHdpdGggIi8iIOKAlCBuZXZlciB0aGUgZnVsbCBVUkwuCi0gUHJlZmVyIGEgdmVyYiBmcm9tIGBlbmRwb2ludHNgIHdoZW4gaXQgZml0czsgb25seSBpbnZlbnQgYSBuZXcgcGF0aCBpZgogIHRoZSB1c2VyIHdhbnRzIHNvbWV0aGluZyBub3QgbGlzdGVkLgotIERlZmF1bHQgdG8gR0VUIHVubGVzcyB0aGUgdXNlciBjbGVhcmx5IGFza2VkIHRvIGNyZWF0ZSAvIHNlbmQgLyB1cGRhdGUgLwogIGRlbGV0ZSBzb21ldGhpbmcuCi0gYHBhcmFtc2AgaXMgZm9yIHF1ZXJ5LXN0cmluZyBhcmdzIChHRVQpOyBgYm9keWAgaXMgZm9yIEpTT04gYm9kaWVzCiAgKFBPU1QvUFVUL1BBVENIKS4gTmV2ZXIgcHV0IGJvZHkgZmllbGRzIHVuZGVyIGBwYXJhbXNgLgotIFVzZSBPTkxZIHRoZSBmaWVsZHMgbGlzdGVkIGluIHRoZSBjaG9zZW4gZW5kcG9pbnQncyBkb2N1bWVudGVkIGBib2R5YCAvCiAgYHBhcmFtc2Agc2NoZW1hLiBORVZFUiBhZGQgZmllbGRzIHRoYXQgYXJlbid0IGluIHRoYXQgc2NoZW1hIOKAlCBtYW55IEFQSXMKICAocGF5bWVudCBnYXRld2F5cyBlc3BlY2lhbGx5KSByZWplY3QgdW5rbm93bi9leHRyYSBmaWVsZHMgd2l0aCBhIGhhcmQgZXJyb3IuCi0gT21pdCBvcHRpb25hbCBmaWVsZHMgeW91IGhhdmUgbm8gcmVhbCB2YWx1ZSBmb3IuIERvbid0IHBhZCB0aGUgYm9keSB3aXRoCiAgZW1wdHksIHBsYWNlaG9sZGVyLCBvciBndWVzc2VkIHZhbHVlcy4KLSBOZXZlciBpbnZlbnQgb3duZXIvcmVwby9jaGFubmVsIGlkcyB0aGUgdXNlciBkaWRuJ3Qgc3VwcGx5IOKAlCBsZWF2ZSB0aGUKICBwYXRoIHBsYWNlaG9sZGVyIGluIGFuZCBzZXQgZW5kcG9pbnQ9bnVsbCB3aXRoIGFuIGV4cGxhbmF0aW9uIGluIHN1bW1hcnkKICBpZiBhIHJlcXVpcmVkIGlkIGlzIG1pc3NpbmcuCg==").decode()

_SUMMARY_EN_SYSTEM = (
    "You are the user-facing voice of an API agent. Given the technical "
    "result of an HTTP call, produce ONE short paragraph (max 3 sentences) "
    "that the end user can read. Plain English, no JSON, no curly braces, "
    "no field names like 'status: 200'. If there's a URL the user should "
    "click, include it as a markdown link — but ONLY use a URL that appears "
    "verbatim in the response data. NEVER invent, guess, or shorten a link. "
    "If the call failed, say what failed and what the user can try next."
)

_SUMMARY_HINGLISH_SYSTEM = (
    "Aap ek API agent ki user-facing voice ho. Technical result ko dekh kar "
    "ek short paragraph (max 3 sentences) Hinglish mein likho jo end user "
    "ko samajh aaye. Plain Hinglish, no JSON, no curly braces, no field "
    "names like 'status: 200'. Agar koi URL hai jo user click kar sakta hai, "
    "use markdown link ke roop mein include karo — par SIRF wahi URL jo "
    "response data mein literally maujood ho. Koi link KABHI invent/guess/short "
    "mat karo. Agar call fail ho gayi, to bolo kya fail hua aur user kya try "
    "kar sakta hai."
)

# Common keys whose value is a user-facing link, best first. Used to pick the
# REAL link to show when the summarizer LLM invents one.
_LINK_KEY_PRIORITY = (
    "short_url", "payment_link", "invoice_url", "payment_url", "url",
    "link", "href", "redirect_url", "checkout_url",
)

def _real_urls_from_response(obj: Any) -> Tuple[List[str], set]:
    """Return (priority_urls, all_urls) actually present in an API response.

    ``priority_urls`` are values of common user-facing link fields (short_url,
    payment_link, …), best first — so we can substitute the RIGHT link when the
    summarizer hallucinated one. Generic: no per-provider knowledge."""
    all_urls: List[str] = []
    priority: List[Tuple[int, str]] = []

    def walk(node: Any) -> None:
        if isinstance(node, dict):
            for k, v in node.items():
                if isinstance(v, str) and v.startswith(("http://", "https://")):
                    all_urls.append(v)
                    lk = k.lower()
                    if lk in _LINK_KEY_PRIORITY:
                        priority.append((_LINK_KEY_PRIORITY.index(lk), v))
                else:
                    walk(v)
        elif isinstance(node, list):
            for it in node:
                walk(it)

    if isinstance(obj, (dict, list)):
        walk(obj)
    elif isinstance(obj, str):
        all_urls.extend(re.findall(r"https?://[^\s\"'<>)\]]+", obj))

    priority.sort(key=lambda t: t[0])
    return [u for _, u in priority], set(all_urls)

def _fix_summary_links(summary: str, response_body: Any) -> str:
    """Replace any URL in the summary that does NOT appear in the real response
    with the actual link from the response (or strip it if there is none).

    Stops the small summarizer model from showing a hallucinated/placeholder
    link (e.g. ``https://rzp.io/i/abc``) instead of the real one. Generic —
    works for any tool that returns a link."""
    if not summary:
        return summary
    priority_urls, real_set = _real_urls_from_response(response_body)
    best = priority_urls[0] if priority_urls else (
        next(iter(real_set)) if real_set else None
    )

    def _md(m: "re.Match") -> str:
        text, url = m.group(1), m.group(2)
        if url in real_set:
            return m.group(0)
        return f"[{text}]({best})" if best else text

    out = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", _md, summary)

    def _bare(m: "re.Match") -> str:
        url = m.group(0)
        if url in real_set:
            return url
        return best or ""

    out = re.sub(r"https?://[^\s\"'<>)\]]+", _bare, out)
    return re.sub(r"[ \t]{2,}", " ", out).strip()

def _strip_json_fences(text: str) -> str:
    text = (text or "").strip()
    if not text:
        return text
    fence = re.match(r"^```(?:json)?\s*([\s\S]*?)\s*```$", text)
    if fence:
        return fence.group(1).strip()
    if text.lower().startswith("json\n"):
        return text[5:].strip()
    return text

def _extract_first_json_object(text: str) -> Optional[str]:
    if not text:
        return None
    text = _strip_json_fences(text)
    depth = 0
    start: Optional[int] = None
    for i, ch in enumerate(text):
        if ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0 and start is not None:
                return text[start : i + 1]
    return None

def _strip_empties(obj: Any) -> Any:
    """Drop None / empty placeholders the local LLM loves to invent."""
    if isinstance(obj, dict):
        cleaned: Dict[str, Any] = {}
        for k, v in obj.items():
            cv = _strip_empties(v)
            if cv is None:
                continue
            if isinstance(cv, str) and cv.strip() == "":
                continue
            if isinstance(cv, (dict, list)) and len(cv) == 0:
                continue
            cleaned[k] = cv
        return cleaned
    if isinstance(obj, list):
        return [
            v for v in (_strip_empties(x) for x in obj)
            if v is not None
            and not (isinstance(v, str) and v.strip() == "")
            and not (isinstance(v, (dict, list)) and len(v) == 0)
        ]
    return obj

# Field names that providers like Razorpay / Stripe / PayPal require to be
# integers. Anything decimal here is a planner mistake we can safely round
# — the LLM has already been told to multiply by 100 (or 1) per the
# provider quirks, so the magnitude is right; only the type is wrong.
_INTEGER_MONEY_KEYS = {
    "amount", "amount_paid", "amount_due", "amount_refunded",
    "unit_amount", "subtotal", "total",
}

def _coerce_integer_money_fields(obj: Any) -> Any:
    """Recursively coerce money fields to int. Handles float 100.0→100 and str '1000'→1000."""
    if isinstance(obj, dict):
        out: Dict[str, Any] = {}
        for k, v in obj.items():
            if k in _INTEGER_MONEY_KEYS and isinstance(v, (float, str)):
                try:
                    out[k] = int(round(float(v)))
                except (ValueError, TypeError):
                    out[k] = v
            elif isinstance(v, (dict, list)):
                out[k] = _coerce_integer_money_fields(v)
            else:
                out[k] = v
        return out
    if isinstance(obj, list):
        return [_coerce_integer_money_fields(x) for x in obj]
    return obj

# Maps LLM-emitted tool name variants to their canonical name. The model
# splits AWS into per-service names ("aws-ec2", "aws-s3", "amazon-web-services")
# even when the system prompt tells them not to — this is the deterministic
# backstop. Add an entry here for every variant we've seen the LLM emit.
_TOOL_ALIASES: Dict[str, str] = {
    # AWS — anything that obviously means Amazon Web Services collapses to "aws"
    "aws-cli": "aws",
    "aws-ec2": "aws",
    "aws-s3": "aws",
    "aws-rds": "aws",
    "aws-lambda": "aws",
    "aws-iam": "aws",
    "aws-sts": "aws",
    "amazon-aws": "aws",
    "amazon-web-services": "aws",
    "amazon-web-service": "aws",
    "amazonaws": "aws",
    "ec2": "aws",
    "s3": "aws",
    # GitHub variants
    "gh": "github",
    "github.com": "github",
    # Gmail variants
    "google-mail": "gmail",
    "googlemail": "gmail",
    # OpenAI variants
    "open-ai": "openai",
    "chatgpt": "openai",
}

def _canonicalize_tool_name(tool_name: Optional[str]) -> Optional[str]:
    """Map LLM-emitted aliases to their canonical seed key. Returns the
    input unchanged when no alias applies."""
    if not tool_name:
        return tool_name
    name = tool_name.strip().lower()
    return _TOOL_ALIASES.get(name, name)

# Per-provider hints for the most common error codes that benefit from a
# concrete next step. Looked up after the response is parsed so the summary
# LLM gets actionable text appended to the body, not just an opaque code.
_SLACK_ERROR_HINTS: Dict[str, str] = {
    "missing_scope": (
        "Slack rejected the call because the bot token is missing one or "
        "more OAuth scopes. To fix: open https://api.slack.com/apps → your "
        "app → OAuth & Permissions → Bot Token Scopes → add the scopes "
        "listed under `needed` in this response → click \"Reinstall to "
        "Workspace\" at the top of the same page (the new scopes don't "
        "take effect until you reinstall). Then paste the fresh "
        "xoxb-… token into the Dynamic Agent."
    ),
    "invalid_auth": (
        "The bot token is wrong or has been revoked. Generate a fresh "
        "Bot User OAuth Token at https://api.slack.com/apps → your app → "
        "OAuth & Permissions."
    ),
    "not_allowed_token_type": (
        "You pasted the wrong KIND of Slack token. The agent needs a "
        "**Bot User OAuth Token** (starts with `xoxb-`), NOT a User OAuth "
        "Token (`xoxp-…`) or App-Level Token (`xapp-…`). "
        "Open https://api.slack.com/apps → your app → OAuth & Permissions. "
        "Under \"OAuth Tokens for Your Workspace\" copy the value labelled "
        "**Bot User OAuth Token** (it's the FIRST one, marked with `xoxb-`). "
        "If you only see a User token, scroll down to \"Bot Token Scopes\" "
        "and add at least one scope (e.g. channels:read) — Slack only "
        "generates a bot token once an app has bot scopes."
    ),
    "token_revoked": (
        "This Slack token has been revoked. Reinstall the app at "
        "https://api.slack.com/apps → your app → \"Install App\" → "
        "\"Reinstall to Workspace\", then paste the fresh xoxb-… token."
    ),
    "token_expired": (
        "This Slack token has expired. Reinstall the app at "
        "https://api.slack.com/apps → your app → \"Install App\" → "
        "\"Reinstall to Workspace\", then paste the fresh xoxb-… token."
    ),
    "account_inactive": (
        "The Slack workspace tied to this token is inactive (paused or "
        "deleted). Use a token from an active workspace."
    ),
    "not_in_channel": (
        "The bot isn't a member of that channel. In Slack, type "
        "`/invite @<your-bot-name>` inside the channel and retry."
    ),
    "channel_not_found": (
        "Slack couldn't find the channel id you provided. Either use the "
        "channel id (starts with C…) or list channels first with "
        "`list slack channels`."
    ),
    "ratelimited": (
        "Slack is rate-limiting this token. Wait 30-60 seconds and retry."
    ),
}

def _interpret_provider_level_error(
    tool: ToolDefinition, http_status: int, parsed_body: Any
) -> int:
    """Catch HTTP-200-with-body-error patterns and remap to a real failure.

    Some providers — Slack is the worst offender — return HTTP 200 even
    when the call failed. Without remapping, the agent reports "success"
    and the user sees only the raw error JSON. We:

      * detect Slack's ``{"ok": false, "error": "<code>"}`` shape and bump
        the status to 400 (or 401/403 for auth-style errors)
      * mutate ``parsed_body`` in place to add a ``hint`` field describing
        the concrete next action the user should take
    """
    if http_status >= 400:
        return http_status

    if not isinstance(parsed_body, dict):
        return http_status

    tool_name = (tool.name or "").lower()

    # Slack-style { ok: false, error: <code> }
    if tool_name == "slack" and parsed_body.get("ok") is False:
        code = (parsed_body.get("error") or "").strip()
        hint = _SLACK_ERROR_HINTS.get(code)
        if hint:
            parsed_body["hint"] = hint
        # Status: auth-style errors → 401, missing scope → 403, else 400.
        if code in ("invalid_auth", "token_revoked", "token_expired"):
            return 401
        if code in ("missing_scope", "no_permission"):
            return 403
        return 400

    return http_status

def _fixup_boto3_kwargs(
    service: str, operation: str, kwargs: Dict[str, Any], region: str
) -> Dict[str, Any]:
    """Rewrite kwargs to match the connection's region / AWS gotchas.

    Small LLMs trip on subtle AWS rules. We don't try to be clever — only
    fix the specific footguns we've seen the planner walk into:

      * S3 create_bucket: us-east-1 rejects CreateBucketConfiguration
        entirely; every other region REQUIRES LocationConstraint to match.
        Whatever the LLM emitted is overridden with the region the user's
        connection is actually configured for.
    """
    if not kwargs:
        kwargs = {}
    out = dict(kwargs)

    if service == "s3" and operation == "create_bucket":
        if (region or "us-east-1") == "us-east-1":
            # us-east-1 is special: passing CreateBucketConfiguration here
            # triggers "InvalidLocationConstraint" no matter what value
            # you put in. The bucket lands in us-east-1 by default.
            out.pop("CreateBucketConfiguration", None)
        else:
            # Force LocationConstraint to the connection region — even if
            # the LLM hallucinated a different one or omitted it.
            out["CreateBucketConfiguration"] = {"LocationConstraint": region}

    return out

def _sanitize_url_string(raw: str) -> str:
    """Strip junk wrappers (angle brackets, quotes, whitespace, trailing
    punctuation) that LLMs love to copy from markdown / man-page snippets.

    ``"<https://api.x.com/v1>"`` → ``"https://api.x.com/v1"``"""
    if not raw:
        return raw
    s = raw.strip()
    # Strip pairs of angle brackets, then quotes — markdown autolinks
    # (`<url>`) and JSON-quoted strings are the common shapes the
    # extractor returns by mistake.
    for opening, closing in [("<", ">"), ('"', '"'), ("'", "'"), ("`", "`")]:
        if s.startswith(opening) and s.endswith(closing) and len(s) >= 2:
            s = s[1:-1].strip()
    # Drop any stray brackets / quotes left mid-string. URLs never contain
    # these characters in practice (they'd need to be percent-encoded).
    s = s.replace("<", "").replace(">", "").replace('"', "").replace("'", "")
    # Trim trailing punctuation that often hitches a ride from prose
    # ("...the base url is https://api.x.com.").
    while s and s[-1] in ".,;:":
        s = s[:-1]
    return s.strip()

# Host fragments that show up in API documentation as *placeholders*, never as
# a real production host. The dynamic extractor sometimes copies one straight
# out of a docs example (e.g. base_url "https://api.example.com"), which would
# send every request into the void. Treat any of these as "no base_url" so the
# fallback chain (LLM training knowledge → real search-result hosts) runs
# instead of persisting a dead host.
_PLACEHOLDER_HOST_MARKERS: Tuple[str, ...] = (
    "example.com", "example.org", "example.net", "example.io", "example.api",
    "api.example", "your-domain", "yourdomain", "your_domain", "your-company",
    "yourcompany", "your-api", "yourapi", "your-app", "api.your", "api.domain",
    "domain.com", "host.com", "hostname", "subdomain", "yourtenant",
    "your-tenant", "mycompany", "myapi", "localhost", "127.0.0.1", "0.0.0.0",
)

def _url_host(url: Optional[str]) -> str:
    """Lower-cased hostname of a URL, without urllib (not imported here)."""
    if not url or not isinstance(url, str):
        return ""
    m = re.match(r"^\s*https?://([^/?#]+)", url, re.IGNORECASE)
    if not m:
        return ""
    return m.group(1).split("@")[-1].split(":")[0].strip().lower()

def _is_placeholder_url(url: Optional[str]) -> bool:
    """True if ``url``'s host looks like a docs placeholder rather than a real
    API host (example.com, your-domain.com, localhost, a leftover {template},
    a reserved test TLD). Used to reject base_urls copied from a docs snippet."""
    if not url or not isinstance(url, str):
        return False
    if "{" in url or "}" in url or "..." in url:  # leftover templating
        return True
    host = _url_host(url)
    if not host:
        return False
    # RFC 2606 / 6761 reserved-for-documentation-and-testing TLDs.
    if host.endswith((".example", ".invalid", ".test", ".local", ".localhost")):
        return True
    return any(marker in host for marker in _PLACEHOLDER_HOST_MARKERS)

_URL_IN_TEXT_RE = re.compile(r"https?://[^\s)>\]}'\"]+", re.IGNORECASE)

def _first_url_in_text(text: Optional[str]) -> Optional[str]:
    """Return the first http(s) URL literally present in ``text`` (e.g. a doc
    link the user pasted into chat), or None. This is plain parsing of what the
    user typed — not URL guessing."""
    if not text or not isinstance(text, str):
        return None
    m = _URL_IN_TEXT_RE.search(text)
    if not m:
        return None
    # Trim trailing punctuation that commonly clings to a pasted URL.
    return m.group(0).rstrip(".,;:!?")

def _is_safe_public_url(url: Optional[str]) -> bool:
    """SSRF guard for USER-SUPPLIED fetch targets (doc/spec links).

    The server fetches these URLs itself, so an attacker could otherwise point
    us at internal services or the cloud metadata endpoint. Allow only http/https
    to a host that resolves entirely to GLOBAL (public) IP addresses. Anything
    private / loopback / link-local (169.254.169.254) / reserved is rejected."""
    import ipaddress
    import socket
    from urllib.parse import urlsplit

    if not url or not isinstance(url, str):
        return False
    try:
        parts = urlsplit(url.strip())
    except Exception:
        return False
    if parts.scheme.lower() not in ("http", "https"):
        return False
    host = parts.hostname
    if not host:
        return False

    def _ip_ok(ip_str: str) -> bool:
        try:
            ip = ipaddress.ip_address(ip_str)
        except ValueError:
            return False
        return not (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_reserved
            or ip.is_multicast
            or ip.is_unspecified
        )

    # Literal IP host — check directly (no DNS).
    try:
        ipaddress.ip_address(host)
        return _ip_ok(host)
    except ValueError:
        pass

    # Hostname — must resolve, and EVERY resolved address must be global.
    try:
        infos = socket.getaddrinfo(host, parts.port or None, proto=socket.IPPROTO_TCP)
    except Exception:
        return False
    addrs = {info[4][0] for info in infos}
    if not addrs:
        return False
    return all(_ip_ok(a) for a in addrs)

# Domains that are never a first-party API reference — tutorials, content
# farms, forums, and third-party "API" resellers. Not hard-dropped (sometimes
# they're the only hit), just pushed to the bottom so the official docs win
# whenever they exist. This is what stops the extractor from learning, say,
# LinkedIn's API from a dev.to / Proxycurl blog.
_LOW_TRUST_DOC_DOMAINS: Tuple[str, ...] = (
    "dev.to", "medium.com", "hashnode.", "freecodecamp.org", "geeksforgeeks.org",
    "towardsdatascience.com", "hackernoon.com", "tutorialspoint.com", "w3schools.com",
    "stackoverflow.com", "stackexchange.com", "reddit.com", "quora.com",
    "youtube.com", "youtu.be", "facebook.com", "twitter.com", "pinterest.",
    "proxycurl.com", "nubela.co", "rapidapi.com", "programmableweb.com",
    "apilist.fun", "public-apis", "wikipedia.org", "blog.", "/blog/",
)

# Path / host fragments that strongly suggest a real API reference or spec.
_OFFICIAL_DOC_URL_SIGNALS: Tuple[str, ...] = (
    "/reference", "/api-reference", "/rest", "/docs/api", "/api/",
    "openapi", "swagger", "/developers", "/developer",
)

def _score_doc_result(result: Dict[str, Any], tool_name: str) -> int:
    """Rank a search result by how likely it is to be the tool's OFFICIAL API
    docs (higher = better). A first-party host (carries the tool's name) + an
    api/developer/docs subdomain + a spec/reference path scores high; content
    farms, forums, and third-party API resellers score negative. This is the
    'human judgment' the small LLM lacks — pick official, ignore blogs."""
    url = (result.get("href") or "").strip().lower()
    if not url:
        return -100
    host = _url_host(url) or url
    score = 0
    # First-party: the host carries the tool's name (linkedin.com, stripe.com…).
    tokens = [t for t in re.split(r"[^a-z0-9]+", (tool_name or "").lower()) if len(t) >= 3]
    if any(t in host for t in tokens):
        score += 50
    # API-surface / docs subdomains.
    if host.startswith(("api.", "developer.", "developers.", "docs.")) or \
            any(s in host for s in (".developer.", ".docs.")):
        score += 20
    # Machine-readable spec — the most reliable source of truth.
    if url.endswith((".json", ".yaml", ".yml")) or "openapi" in url or "swagger" in url:
        score += 25
    # Reference-y path segments.
    if any(seg in url for seg in _OFFICIAL_DOC_URL_SIGNALS):
        score += 12
    # Content farms / forums / third-party resellers — push to the bottom.
    if any(d in host or d in url for d in _LOW_TRUST_DOC_DOMAINS):
        score -= 80
    return score

def _registrable_domain(host: str) -> str:
    """Last two labels of a host — a cheap 'registrable domain' proxy
    (api.linkedin.com → linkedin.com). Good enough to compare two hosts."""
    parts = (host or "").split(".")
    return ".".join(parts[-2:]) if len(parts) >= 2 else (host or "")

def _host_relates_to_tool(
    base_url: Optional[str], tool_name: str, docs_url: Optional[str] = None
) -> bool:
    """Trust check: does ``base_url``'s host plausibly belong to this tool?

    True if the host carries a tool-name token (api.linkedin.com ← 'linkedin'),
    or shares a registrable domain with the official docs URL. This catches the
    small model hallucinating a real-but-WRONG host — e.g. returning
    base_url=api.github.com for the tool 'linkedin'. When it returns False we
    distrust the web base_url and prefer the model's own knowledge.

    Conservative by design: a tool whose API host genuinely differs from its
    name (gmail → googleapis.com) returns False here, which only costs a
    fallback to LLM knowledge — never a wrong save."""
    host = _url_host(base_url)
    if not host:
        return False
    tokens = [t for t in re.split(r"[^a-z0-9]+", (tool_name or "").lower()) if len(t) >= 4]
    if any(t in host for t in tokens):
        return True
    dhost = _url_host(docs_url or "")
    if dhost and _registrable_domain(host) == _registrable_domain(dhost):
        return True
    return False

def _extract_text_from_bytes(
    data: bytes, name: str = "", content_type: str = ""
) -> str:
    """Turn an arbitrary document (uploaded OR downloaded) into clean text the
    LLM can read — whatever the format. End users are dumb: they'll send a PDF,
    a Word doc, an HTML page, a screenshot-export, anything. We sniff the type
    by magic bytes + filename + content-type and extract readable text.

    Handles: PDF, Word (.docx), HTML, JSON/YAML, Markdown, plain text. Unknown
    binary falls back to best-effort UTF-8. (Image OCR is out of scope here —
    that needs a vision model / tesseract; flagged separately.)"""
    if not data:
        return ""
    nm = (name or "").lower()
    ct = (content_type or "").lower()
    head = data[:8]

    # --- PDF ---
    if head.startswith(b"%PDF") or nm.endswith(".pdf") or "application/pdf" in ct:
        try:
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception as e:
            logger.warning(f"PDF text extraction failed: {e}")
            return ""

    # --- Word .docx (zip: 'PK' magic + .docx name / mime) ---
    if (head.startswith(b"PK") and nm.endswith(".docx")) or \
            "officedocument.wordprocessingml" in ct:
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:  # docs often put endpoints in tables
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        except Exception as e:
            logger.warning(f"DOCX text extraction failed: {e}")
            return ""

    # --- everything else: decode then maybe clean ---
    try:
        text = data.decode("utf-8", "ignore")
    except Exception:
        return ""
    stripped = text.lstrip()

    # JSON / YAML / spec text → return raw so the OpenAPI parser can try it.
    if nm.endswith((".json", ".yaml", ".yml")) or "json" in ct or "yaml" in ct \
            or stripped[:1] == "{" or stripped[:8] == "openapi:" or stripped[:8] == "swagger:":
        return text

    # HTML → strip tags to readable text (trafilatura, then BeautifulSoup).
    if nm.endswith((".html", ".htm")) or "text/html" in ct \
            or "<html" in stripped[:300].lower() or "<!doctype html" in stripped[:300].lower():
        try:
            import trafilatura
            cleaned = trafilatura.extract(text) or ""
            if cleaned.strip():
                return cleaned
        except Exception:
            pass
        try:
            from bs4 import BeautifulSoup
            return BeautifulSoup(text, "html.parser").get_text("\n")
        except Exception:
            return text

    # plain text / markdown / unknown — already decoded.
    return text

def _has_api_signals(text: str) -> bool:
    """Cheap check: does this text plausibly document API endpoints, vs. being
    a pure marketing/index page? Kept lenient — it only decides whether a
    crawled page CONTRIBUTES its text to the extractor, so false-positives just
    give the LLM a bit more to read, while false-negatives silently drop real
    endpoint pages (the worse failure — that's how a whole doc set collapsed to
    one endpoint)."""
    if not text:
        return False
    low = text.lower()
    verbs = sum(low.count(v) for v in ("get ", "post ", "put ", "patch ", "delete "))
    paths = len(re.findall(r"/[a-z0-9_]+(?:/[a-z0-9_{}]+)+", low))
    # Endpoint-ish vocabulary common on REST reference pages that describe calls
    # in prose/tables rather than as clean "GET /path" lines (e.g. LinkedIn on
    # Microsoft Learn).
    kw = sum(low.count(k) for k in (
        "endpoint", "request", "response", "api.", "/v1/", "/v2/",
        "parameter", "header", "scope", "https://api",
    ))
    return (verbs >= 1 and paths >= 1) or paths >= 3 or kw >= 4

def _render_html(url: str, timeout_ms: int = 20000, settle_ms: int = 1800) -> Optional[str]:
    """Render a URL in headless Chromium and return the post-JavaScript HTML.

    For JS-heavy doc sites (Microsoft Learn, many modern API portals) the
    static HTML a plain GET returns is an empty shell — the real content and
    nav links are injected by JavaScript. A headless browser executes that JS
    server-side (invisible, no GUI) so we can read what a human would see.

    Returns None if Playwright/Chromium isn't available or the render fails —
    callers fall back to the plain-HTTP text."""
    try:
        from playwright.sync_api import sync_playwright
    except Exception:
        logger.debug("playwright not installed; skipping headless render")
        return None
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                args=["--no-sandbox", "--disable-dev-shm-usage"]
            )
            try:
                page = browser.new_page(
                    user_agent="Mozilla/5.0 (compatible; Adaptora-DocBot/1.0)"
                )
                page.goto(url, timeout=timeout_ms, wait_until="domcontentloaded")
                try:
                    page.wait_for_load_state("networkidle", timeout=6000)
                except Exception:
                    pass
                page.wait_for_timeout(settle_ms)
                return page.content()
            finally:
                browser.close()
    except Exception as e:
        logger.warning(f"headless render failed for {url}: {e}")
        return None

def _normalize_endpoint(endpoint: str) -> str:
    """Coerce whatever the LLM produced into a relative path under base_url."""
    if not endpoint:
        return endpoint
    e = _sanitize_url_string(endpoint)
    if e.lower().startswith(("http://", "https://")):
        from urllib.parse import urlsplit

        parts = urlsplit(e)
        e = parts.path or "/"
        if parts.query:
            e = f"{e}?{parts.query}"
    if not e.startswith("/"):
        e = "/" + e
    return e

def _slug_from_path(method: str, path: str) -> str:
    """Build a stable verb slug from an HTTP method + path. Used as a
    fallback when an OpenAPI operation has no operationId. ``GET /repos/{owner}/{repo}/issues``
    → ``get_repos_owner_repo_issues``. Path placeholders keep their bare
    name so different param positions don't collide."""
    tail = re.sub(r"[{}]", "", path or "")
    tail = re.sub(r"[^a-zA-Z0-9]+", "_", tail).strip("_")
    return f"{method.lower()}_{tail}" if tail else method.lower()

def _resolve_json_pointer(root: Any, pointer: str) -> Any:
    """Resolve an RFC-6901 JSON Pointer (``#/paths/~1lists``) against
    ``root``. Used to expand ``$ref`` references inside OpenAPI/Swagger
    specs — Mailchimp's spec is the motivating example: each path is a
    ``{"$ref": "#/paths/~1lists"}`` pointer to the actual operations
    block. Returns the resolved node or None on any error (bad pointer,
    missing key, type mismatch). Only handles same-document refs (which
    is what the OpenAPI spec actually allows at the paths object)."""
    if not isinstance(pointer, str) or not pointer.startswith("#"):
        return None
    parts = pointer[1:].lstrip("/").split("/")
    if parts == [""]:
        return root
    node: Any = root
    for raw in parts:
        # RFC 6901 escapes: ~1 → /, ~0 → ~ (in that order to avoid
        # double-decoding accidentally re-introduced ~ characters).
        token = raw.replace("~1", "/").replace("~0", "~")
        if isinstance(node, dict) and token in node:
            node = node[token]
        elif isinstance(node, list):
            try:
                node = node[int(token)]
            except (ValueError, IndexError):
                return None
        else:
            return None
    return node

def _openapi_collect_schema_props(
    schema: Any, spec: Dict[str, Any], out: Dict[str, str], depth: int = 0
) -> None:
    """Walk an OpenAPI request-body schema and collect its top-level property
    names into ``out`` as ``{field: description}``. Resolves ``$ref`` and folds
    ``allOf``/``oneOf``/``anyOf`` so refs-to-component-schemas (the common case)
    actually yield their fields. Bounded depth so a cyclic spec can't loop."""
    if not isinstance(schema, dict) or depth > 6:
        return
    if "$ref" in schema and isinstance(schema["$ref"], str):
        schema = _resolve_json_pointer(spec, schema["$ref"]) or {}
    for comb in ("allOf", "oneOf", "anyOf"):
        for sub in schema.get(comb) or []:
            _openapi_collect_schema_props(sub, spec, out, depth + 1)
    # If it's an object with properties, harvest the field names.
    props = schema.get("properties")
    if isinstance(props, dict):
        required = set(schema.get("required") or [])
        for name, p in props.items():
            if not isinstance(name, str) or name in out:
                continue
            desc = ""
            if isinstance(p, dict):
                desc = (p.get("description") or p.get("type") or "").strip()
            if name in required:
                desc = (desc + " (required)").strip()
            out[name] = desc or "field"

def _openapi_body_fields(op: Dict[str, Any], spec: Dict[str, Any]) -> Optional[Dict[str, str]]:
    """Extract the JSON request-body fields of an OpenAPI operation as
    ``{field: description}`` (or None). This is what lets the planner/grounding
    know which body fields a write endpoint actually accepts — without it every
    spec-derived POST had an empty body and invalid fields slipped through."""
    if not isinstance(op, dict):
        return None
    rb = op.get("requestBody")
    if isinstance(rb, dict) and "$ref" in rb and isinstance(rb["$ref"], str):
        rb = _resolve_json_pointer(spec, rb["$ref"])
    if not isinstance(rb, dict):
        return None
    content = rb.get("content")
    if not isinstance(content, dict):
        return None
    media = content.get("application/json") or next(
        (v for v in content.values() if isinstance(v, dict)), None
    )
    if not isinstance(media, dict):
        return None
    fields: Dict[str, str] = {}
    _openapi_collect_schema_props(media.get("schema"), spec, fields)
    return fields or None

def _oauth1_auth_header(
    method: str,
    url: str,
    query_params: Optional[Dict[str, Any]],
    *,
    consumer_key: str,
    consumer_secret: str,
    token: str = "",
    token_secret: str = "",
) -> str:
    """Build an OAuth 1.0a ``Authorization: OAuth …`` header (HMAC-SHA1).

    Implemented with the stdlib only — no oauthlib dependency — so OAuth1
    providers (Twitter v1.1, Trello, Tumblr, Garmin, …) accept the call. Signs
    the request method + base URI + sorted oauth_*/query params per RFC 5849.
    Body params are intentionally excluded: we only send JSON/no bodies, and
    OAuth1 only folds *form-encoded* bodies into the signature base."""
    import base64 as _b64
    import hmac as _hmac
    import hashlib as _hashlib
    import secrets as _secrets
    import time as _t
    from urllib.parse import urlsplit, urlunsplit, parse_qsl, quote

    def _q(s: Any) -> str:
        # RFC 3986 percent-encoding (unreserved chars kept, incl. '~').
        return quote(str(s), safe="~")

    split = urlsplit(url)
    base_uri = urlunsplit(
        (split.scheme.lower(), split.netloc.lower(), split.path, "", "")
    )

    oauth_params: Dict[str, str] = {
        "oauth_consumer_key": consumer_key,
        "oauth_nonce": _secrets.token_hex(16),
        "oauth_signature_method": "HMAC-SHA1",
        "oauth_timestamp": str(int(_t.time())),
        "oauth_version": "1.0",
    }
    if token:
        oauth_params["oauth_token"] = token

    # Everything that participates in the signature: oauth_* + query string +
    # any caller-supplied query params, all percent-encoded then sorted.
    sig_params: List[Tuple[str, str]] = [(_q(k), _q(v)) for k, v in oauth_params.items()]
    for k, v in (query_params or {}).items():
        sig_params.append((_q(k), _q(v)))
    for k, v in parse_qsl(split.query, keep_blank_values=True):
        sig_params.append((_q(k), _q(v)))
    sig_params.sort()
    param_string = "&".join(f"{k}={v}" for k, v in sig_params)

    base_string = "&".join([method.upper(), _q(base_uri), _q(param_string)])
    signing_key = f"{_q(consumer_secret)}&{_q(token_secret)}"
    signature = _b64.b64encode(
        _hmac.new(signing_key.encode(), base_string.encode(), _hashlib.sha1).digest()
    ).decode()
    oauth_params["oauth_signature"] = signature

    return "OAuth " + ", ".join(
        f'{_q(k)}="{_q(v)}"' for k, v in sorted(oauth_params.items())
    )

def _ollama_unreachable_hint(exc: Exception) -> str:
    """Actionable error for the most common Ollama-unreachable scenarios.

    A bare `ConnectTimeoutError` exception string scares users into thinking
    something complex is broken when it's almost always one of three boring
    deployment issues. List them inline so the user can fix it without
    grepping logs."""
    return (
        f"Cannot reach Ollama at {settings.OLLAMA_API_URL}: "
        f"{exc.__class__.__name__}. Most likely one of:\n"
        "  1. Wrong port — Ollama listens on 11434 by default, not 80.\n"
        "  2. EC2 security group blocks inbound TCP on the Ollama port "
        "from this backend's IP.\n"
        "  3. Ollama bound to 127.0.0.1 only — set "
        "OLLAMA_HOST=0.0.0.0:11434 on the EC2 and restart Ollama.\n"
        "Smoke test from this server: "
        f"`curl --max-time 5 {settings.OLLAMA_API_URL}/api/tags`."
    )

def _ollama_chat_json(
    system: str,
    user: str,
    *,
    temperature: float = 0.0,
    num_predict: int = 512,
    num_ctx: int = 4096,
) -> Dict[str, Any]:
    """Round-trip to Ollama /api/chat with JSON-mode hinting."""
    if not settings.OLLAMA_API_URL or not settings.OLLAMA_MODEL:
        raise RuntimeError(
            "Ollama is not configured (OLLAMA_API_URL / OLLAMA_MODEL). "
            "The dynamic agent needs a local model to reason."
        )

    try:
        resp = requests.post(
            f"{settings.OLLAMA_API_URL}/api/chat",
            json={
                "model": settings.OLLAMA_MODEL,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                "format": "json",
                "stream": False,
                # Pin the model in memory so back-to-back calls in a single
                # run_turn (and across turns) don't pay the ~30s cold reload.
                "keep_alive": settings.OLLAMA_KEEP_ALIVE,
                "options": {
                    "temperature": temperature,
                    "num_predict": num_predict,
                    "num_ctx": num_ctx,
                },
            },
            timeout=(settings.OLLAMA_CONNECT_TIMEOUT, settings.OLLAMA_TIMEOUT),
        )
    except (requests.ConnectTimeout, requests.ConnectionError) as exc:
        # TCP-level failure (timeout, refused, name resolution). Surface
        # the deployment checklist instead of the raw urllib3 stack.
        raise RuntimeError(_ollama_unreachable_hint(exc)) from exc
    resp.raise_for_status()
    data = resp.json()
    text = (data.get("message") or {}).get("content", "") or data.get("response", "")
    if not text:
        raise RuntimeError("Empty response from Ollama")

    candidate = _extract_first_json_object(text) or _strip_json_fences(text)
    try:
        return json.loads(candidate)
    except json.JSONDecodeError as exc:
        raise RuntimeError(
            f"LLM did not return parseable JSON. Raw: {text!r}"
        ) from exc

def _ollama_chat_text(
    system: str,
    user: str,
    *,
    temperature: float = 0.2,
    num_predict: int = 256,
    num_ctx: int = 4096,
) -> str:
    """Plain-text chat completion (used for the user-facing summary)."""
    if not settings.OLLAMA_API_URL or not settings.OLLAMA_MODEL:
        return ""
    try:
        resp = requests.post(
            f"{settings.OLLAMA_API_URL}/api/chat",
            json={
                "model": settings.OLLAMA_MODEL,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                "stream": False,
                # Pin the model in memory (see _ollama_chat_json) so the
                # summary call doesn't trigger a cold reload either.
                "keep_alive": settings.OLLAMA_KEEP_ALIVE,
                "options": {
                    "temperature": temperature,
                    "num_predict": num_predict,
                    "num_ctx": num_ctx,
                },
            },
            timeout=(settings.OLLAMA_CONNECT_TIMEOUT, settings.OLLAMA_TIMEOUT),
        )
        resp.raise_for_status()
        data = resp.json()
        return (data.get("message") or {}).get("content", "") or data.get("response", "") or ""
    except (requests.ConnectTimeout, requests.ConnectionError) as exc:
        # Surface the deployment checklist so the user knows it's a
        # networking issue, not a model / parsing issue.
        logger.warning(_ollama_unreachable_hint(exc))
        return ""
    except Exception as exc:
        logger.warning(f"summary LLM call failed: {exc}")
        return ""

def warmup_model(timeout: float = 60.0) -> bool:
    """Load the model into Ollama's memory ahead of the first real request.

    The first inference after the model is unloaded pays a ~30s cold load.
    Firing a tiny generation at startup (and pinning it with keep_alive)
    means the user's first `run_action` is already warm instead of eating
    that reload inside the request — which is exactly what tripped the MCP
    client's timeout. Best-effort: never raises, just logs and returns a
    bool so the caller can decide whether to retry."""
    if not settings.OLLAMA_API_URL or not settings.OLLAMA_MODEL:
        return False
    try:
        resp = requests.post(
            f"{settings.OLLAMA_API_URL}/api/chat",
            json={
                "model": settings.OLLAMA_MODEL,
                "messages": [{"role": "user", "content": "ok"}],
                "stream": False,
                "keep_alive": settings.OLLAMA_KEEP_ALIVE,
                # num_ctx MUST match the value the real calls use (the 4096
                # default in _ollama_chat_json / _ollama_chat_text). If they
                # differ, the first real request re-loads the model with a
                # different context size and the warmup is wasted.
                "options": {"num_predict": 1, "num_ctx": 4096},
            },
            timeout=(settings.OLLAMA_CONNECT_TIMEOUT, timeout),
        )
        resp.raise_for_status()
        logger.info(
            f"Ollama model '{settings.OLLAMA_MODEL}' warmed up and pinned "
            f"(keep_alive={settings.OLLAMA_KEEP_ALIVE})."
        )
        return True
    except Exception as exc:
        logger.warning(f"Ollama warmup failed (will load on first request): {exc}")
        return False

class DynamicAgentError(Exception):
    """Raised for caller-visible errors so route handlers can map to HTTP."""

class DynamicAgentService:
    """End-to-end orchestrator: identify → docs → auth → plan → execute."""

    def __init__(self, llm: Optional[LLMProvider] = None):
        self.llm = llm or LLMProvider()

    # =================================================== step 1: identify

    def identify_tool(self, prompt: str) -> Dict[str, Any]:
        """LLM-only tool identification. No catalog, no shortlist."""
        try:
            decision = _ollama_chat_json(
                _TOOL_IDENTIFY_SYSTEM,
                f"User prompt: {prompt!r}\n\nReturn the JSON envelope.",
                temperature=0.0,
                # Output is a tiny 5-field JSON, so cap generation short.
                # IMPORTANT: keep num_ctx at the shared default — changing
                # num_ctx between calls forces Ollama to reload the model
                # (a ~30s cold load), which would defeat keep_alive. Every
                # call in a turn must use the SAME num_ctx to stay warm.
                num_predict=128,
            )
        except Exception as exc:
            logger.warning(f"identify_tool: LLM failed: {exc}")
            return {
                "tool": None,
                "intent": "ambiguous",
                "wants_doc_import": False,
                "confidence": 0.0,
                "reason": f"LLM error: {exc}",
            }

        tool = (decision.get("tool") or "").strip().lower() or None
        tool = _canonicalize_tool_name(tool)
        intent = decision.get("intent") or "ambiguous"
        if intent not in {"connect", "action", "ambiguous"}:
            intent = "ambiguous"
        try:
            confidence = float(decision.get("confidence") or 0.0)
        except (TypeError, ValueError):
            confidence = 0.0
        return {
            "tool": tool,
            "intent": intent,
            "confidence": confidence,
            "reason": decision.get("reason") or "",
        }

    # =================================================== step 2: docs

    def lookup_or_fetch_docs(
        self,
        db: Session,
        tool_name: str,
        *,
        force_refresh: bool = False,
        status_callback: Optional[Callable[[str, Dict[str, Any]], None]] = None,
    ) -> Optional[ToolDefinition]:
        """DB cache first, then web search + LLM extract from fetched docs.

        Everything is derived from documents we actually fetch: there are no
        seeds, no templated host guesses, and no LLM training-data fallback.
        If the web search can't surface a usable doc/spec for the tool, this
        returns None.

        ``status_callback(step, data)`` is called at each major stage so the
        streaming refresh endpoint can surface progress to the user."""

        def _emit(step: str, data: Optional[Dict[str, Any]] = None) -> None:
            if status_callback is not None:
                try:
                    status_callback(step, data or {})
                except Exception:
                    logger.exception("status_callback raised; ignoring")

        row = (
            db.query(ToolDefinition)
            .filter(ToolDefinition.name == tool_name)
            .first()
        )
        _emit("starting", {"tool": tool_name, "force_refresh": force_refresh})

        if row and not force_refresh:
            _emit("cache_hit", {"source": row.source})
            return row

        # Web search + LLM extract — runs on first load and on force_refresh.
        try:
            extracted = self._extract_docs_from_web(
                tool_name,
                status_callback=status_callback,
            )
        except Exception as exc:
            logger.exception(f"docs extraction failed for {tool_name}")
            _emit("error", {"reason": f"extraction failed: {exc}"})
            extracted = None

        # No fallbacks: if web extraction couldn't surface a usable base_url
        # from real fetched docs, there's nothing actionable to save.
        if not extracted or not extracted.get("base_url"):
            _emit("error", {"reason": "no usable docs found on the web for this tool"})
            return None

        source = "web"
        display_name = extracted.get("display_name") or tool_name.title()

        if row:
            row.display_name = display_name
            row.base_url = extracted["base_url"]
            row.auth_type = extracted.get("auth_type") or "API_KEY"
            row.auth_config = extracted.get("auth_config") or {}
            row.endpoints = extracted.get("endpoints") or {}
            row.rate_limits = extracted.get("rate_limits")
            row.examples = extracted.get("examples")
            row.quirks = extracted.get("quirks")
            row.docs_url = extracted.get("docs_url")
            row.source = source
            row.last_fetched_at = datetime.utcnow()
        else:
            row = ToolDefinition(
                name=tool_name,
                display_name=display_name,
                base_url=extracted["base_url"],
                auth_type=extracted.get("auth_type") or "API_KEY",
                auth_config=extracted.get("auth_config") or {},
                endpoints=extracted.get("endpoints") or {},
                rate_limits=extracted.get("rate_limits"),
                examples=extracted.get("examples"),
                quirks=extracted.get("quirks"),
                docs_url=extracted.get("docs_url"),
                source=source,
            )
            db.add(row)
        try:
            db.commit()
        except StaleDataError:
            # The row was deleted/changed by a concurrent request between our
            # load and this UPDATE (e.g. two refreshes racing). Don't crash the
            # refresh — roll back, then re-insert the freshly-extracted data as
            # a new row keyed by name.
            logger.warning(
                f"stale row during {tool_name} save; rolling back and re-inserting"
            )
            db.rollback()
            db.query(ToolDefinition).filter(
                ToolDefinition.name == tool_name
            ).delete()
            db.commit()
            row = ToolDefinition(
                name=tool_name,
                display_name=display_name,
                base_url=extracted["base_url"],
                auth_type=extracted.get("auth_type") or "API_KEY",
                auth_config=extracted.get("auth_config") or {},
                endpoints=extracted.get("endpoints") or {},
                rate_limits=extracted.get("rate_limits"),
                examples=extracted.get("examples"),
                quirks=extracted.get("quirks"),
                docs_url=extracted.get("docs_url"),
                source=source,
            )
            db.add(row)
            db.commit()
        db.refresh(row)
        _emit(
            "saved",
            {
                "endpoint_count": len(row.endpoints or {}),
                "auth_type": row.auth_type,
                "source": source,
                "has_rate_limits": row.rate_limits is not None,
                "examples_count": len(row.examples or []),
            },
        )
        return row

    def _render_and_crawl(
        self,
        url: str,
        status_callback: Optional[Callable[[str, Dict[str, Any]], None]] = None,
        max_pages: int = 40,
        max_depth: int = 3,
        time_budget_s: float = 300.0,
        char_cap: int = 220000,
    ) -> str:
        """Give it an INDEX/landing URL; it recursively renders and crawls the
        whole same-section doc tree (headless, server-side) and returns the
        aggregated readable text of every page that actually lists endpoints.

        Breadth-first within the URL's path section (e.g. /en-us/linkedin/*),
        rendering each page so JavaScript-built nav + content are visible. Index
        pages contribute their links; pages with real API signals contribute
        their text. Bounded by max_pages / max_depth / a wall-clock budget so a
        giant doc site can't run forever. The start page's text is always kept
        (it usually holds the base URL + auth overview)."""
        from urllib.parse import urlsplit
        from bs4 import BeautifulSoup
        import time as _time

        def _emit(step: str, data: Optional[Dict[str, Any]] = None) -> None:
            if status_callback:
                try:
                    status_callback(step, data or {})
                except Exception:
                    pass

        base = urlsplit(url)
        section = base.path.rstrip("/") or "/"
        deadline = _time.time() + time_budget_s
        visited: set = set()
        queue: List[Tuple[str, int]] = [(url, 0)]
        collected: List[str] = []
        total_chars = 0
        rendered = 0

        while queue and rendered < max_pages and _time.time() < deadline:
            cur, depth = queue.pop(0)
            if cur in visited:
                continue
            visited.add(cur)
            html = _render_html(cur, settle_ms=1200)
            rendered += 1
            if not html:
                continue
            text = BeautifulSoup(html, "html.parser").get_text("\n")
            # Keep text from the start page (base/auth context) and from any
            # page that actually documents endpoints.
            if cur == url or _has_api_signals(text):
                snippet = f"### {cur}\n{text}"
                collected.append(snippet)
                total_chars += len(snippet)
                _emit("page_collected", {"url": cur, "pages": len(collected)})
            # Harvest deeper same-section links from the rendered DOM.
            if depth < max_depth and total_chars < char_cap:
                links = self._extract_links_from_html(html, cur, max_links=200)
                cands: List[str] = []
                for l in links:
                    ls = urlsplit(l)
                    if (
                        ls.netloc == base.netloc
                        and ls.path.rstrip("/").startswith(section)
                        and l not in visited
                    ):
                        cands.append(l)
                # Most API-reference-looking links first.
                for l in sorted(set(cands), key=self._score_api_link, reverse=True):
                    queue.append((l, depth + 1))
            _emit("crawl_progress", {
                "rendered": rendered, "queued": len(queue),
                "collected": len(collected),
            })
            if total_chars >= char_cap:
                break

        return "\n\n".join(collected)

    def _resolve_paid_llm(
        self, db: Optional[Session], user_id: Optional[int]
    ) -> Optional[Tuple[str, str, str]]:
        """Find a usable paid model for high-accuracy extraction.

        Prefers a per-user key added on the dashboard (UserAPIKey), then the
        env-configured keys. Returns (provider, model, api_key) or None. Claude
        is preferred over GPT for structured-doc extraction."""
        # 1) Per-user dashboard keys (encrypted).
        if db is not None and user_id is not None:
            try:
                rows = (
                    db.query(UserAPIKey)
                    .filter(UserAPIKey.user_id == user_id, UserAPIKey.is_active == True)  # noqa: E712
                    .all()
                )
                by_provider = {}
                for r in rows:
                    prov = (r.provider or "").lower()
                    if prov in ("anthropic", "openai") and r.api_key:
                        by_provider.setdefault(prov, r)
                for prov in ("anthropic", "openai"):
                    r = by_provider.get(prov)
                    if r:
                        try:
                            key = decrypt_api_key(r.api_key)
                        except Exception:
                            key = r.api_key  # tolerate plaintext legacy rows
                        if key:
                            default_model = (
                                settings.ANTHROPIC_MODEL if prov == "anthropic"
                                else settings.OPENAI_MODEL
                            )
                            return prov, (r.model_name or default_model), key
            except Exception as exc:
                logger.debug(f"per-user key lookup failed: {exc}")
        # 2) Env-configured keys.
        if getattr(settings, "ANTHROPIC_API_KEY", None):
            return "anthropic", settings.ANTHROPIC_MODEL, settings.ANTHROPIC_API_KEY
        if getattr(settings, "OPENAI_API_KEY", None):
            return "openai", settings.OPENAI_MODEL, settings.OPENAI_API_KEY
        return None

    def _extract_json_smart(
        self,
        system: str,
        user: str,
        *,
        num_predict: int = 8192,
        num_ctx: int = settings.OLLAMA_EXTRACT_NUM_CTX,
        db: Optional[Session] = None,
        user_id: Optional[int] = None,
        status_callback: Optional[Callable[[str, Dict[str, Any]], None]] = None,
    ) -> Dict[str, Any]:
        """Extract a JSON envelope, preferring a configured paid model (Claude/
        GPT) for accuracy and falling back to the local LLM otherwise.

        ``num_ctx`` only affects the local fallback. Doc extraction wants a big
        window (32768) for large doc chunks, but callers on the hot path (e.g.
        planning) MUST pass the shared 4096 default — a different num_ctx forces
        Ollama to reload the model, which stalls back-to-back calls.

        When no paid key is available we emit a one-time ``model_warning`` so
        the UI can tell the user 'using the local model — add a Claude/OpenAI
        key for higher accuracy on complex docs'."""
        def _emit(step: str, data: Optional[Dict[str, Any]] = None) -> None:
            if status_callback:
                try:
                    status_callback(step, data or {})
                except Exception:
                    pass

        paid = self._resolve_paid_llm(db, user_id)
        if paid:
            provider, model, key = paid
            prompt = (
                f"{system}\n\n{user}\n\n"
                "Respond with ONLY the JSON object — no prose, no markdown fences."
            )
            try:
                if provider == "anthropic":
                    text, _ = self.llm.query_anthropic(
                        prompt, model=model, temperature=0.0, api_key=key,
                        max_tokens=min(num_predict, 8192),
                    )
                else:
                    text, _ = self.llm.query_openai(
                        prompt, model=model, temperature=0.0, api_key=key,
                        max_tokens=min(num_predict, 8192),
                    )
                blob = _extract_first_json_object(text) or text
                return json.loads(blob)
            except Exception as exc:
                logger.warning(
                    f"paid-model extraction failed ({provider}); "
                    f"falling back to local model: {exc}"
                )
                _emit("model_warning", {
                    "reason": f"{provider} extraction failed, using local model",
                })
                # fall through to Ollama
        else:
            _emit("model_warning", {
                "reason": "no Claude/OpenAI key configured — using the local "
                          "model. Add a key on the dashboard for higher accuracy "
                          "on complex docs.",
            })
        return _ollama_chat_json(
            system, user, temperature=0.0, num_predict=num_predict, num_ctx=num_ctx
        )

    @staticmethod
    def _merge_endpoint_def(existing: Any, new: Any) -> Dict[str, Any]:
        """Merge two extracted definitions of the SAME endpoint, keeping the
        richer of each part. Unions body/params field lists so a later chunk
        that documents the request body isn't discarded just because an earlier
        chunk mentioned the endpoint first."""
        if not isinstance(existing, dict):
            return new if isinstance(new, dict) else {}
        if not isinstance(new, dict):
            return existing
        out = dict(existing)
        for key in ("method", "path", "description"):
            if not out.get(key) and new.get(key):
                out[key] = new[key]
        for key in ("body", "params"):
            ev, nv = out.get(key), new.get(key)
            if isinstance(ev, dict) and isinstance(nv, dict):
                out[key] = {**nv, **ev}  # union; existing descriptions win
            elif not ev and nv:
                out[key] = nv
        return out

    def _extract_credential_labels(
        self,
        tool_name: str,
        auth_type: str,
        doc_snippet: str,
        db: Optional[Session] = None,
        user_id: Optional[int] = None,
    ) -> Dict[str, Any]:
        """Focused second pass to get the provider's REAL credential names.

        The big extraction prompt often skips credential_field_overrides on a
        small local model, leaving the UI showing generic 'Username/Password'.
        This narrow question ("what does this provider call its credentials?")
        is something even a 7B answers reliably. Returns a
        credential_field_overrides dict (possibly empty)."""
        at = (auth_type or "").upper()
        if at == "BASIC":
            fields, shape = (
                '"username" and "password"',
                '{"username": {"label": "..."}, "password": {"label": "..."}}',
            )
        elif at in ("API_KEY", "BEARER", "PAT"):
            fields, shape = '"secret"', '{"secret": {"label": "..."}}'
        else:
            return {}
        system = (
            "You name API credential fields. Given an API and a doc snippet, "
            "return the EXACT human-facing names the provider uses for its "
            "credentials (as shown in its dashboard) so a form doesn't show a "
            "generic 'Username/Password'. Return STRICT JSON only, no prose."
        )
        user = (
            f"API: {tool_name}\nAuth type: {at}\nFields to label: {fields}\n\n"
            f"DOC SNIPPET:\n{(doc_snippet or '')[:6000]}\n\n"
            f"Return JSON of shape {shape}. Use the provider's real names "
            f"(examples: Razorpay -> Key ID / Key Secret; Twilio -> Account SID "
            f"/ Auth Token; Stripe -> Secret Key). If the docs genuinely use the "
            f"literal words username/password, return {{}}."
        )
        try:
            out = self._extract_json_smart(
                # Same window as doc extraction so this follow-up call doesn't
                # force an Ollama model reload mid-import.
                system, user, num_predict=128,
                num_ctx=settings.OLLAMA_EXTRACT_NUM_CTX, db=db, user_id=user_id,
            )
        except Exception:
            return {}
        result: Dict[str, Any] = {}
        if isinstance(out, dict):
            for fname in ("username", "password", "secret"):
                v = out.get(fname)
                if isinstance(v, dict) and isinstance(v.get("label"), str) and v["label"].strip():
                    result[fname] = {"label": v["label"].strip()}
        return result

    def _apply_credential_labels(
        self,
        extracted: Dict[str, Any],
        tool_name: str,
        doc_text: str,
        db: Optional[Session] = None,
        user_id: Optional[int] = None,
    ) -> None:
        """Ensure friendly credential labels exist. No-op if the main extraction
        already set them, or the auth type doesn't take a pasted secret."""
        at = (extracted.get("auth_type") or "").upper()
        if at not in ("BASIC", "API_KEY", "BEARER", "PAT"):
            return
        cfg = extracted.get("auth_config") or {}
        if cfg.get("credential_field_overrides"):
            return
        labels = self._extract_credential_labels(
            tool_name, at, doc_text, db=db, user_id=user_id
        )
        if labels:
            cfg["credential_field_overrides"] = labels
            extracted["auth_config"] = cfg

    def _extract_from_doc_text(
        self,
        text: str,
        tool_name: str,
        origin: str,
        db: Optional[Session] = None,
        user_id: Optional[int] = None,
        status_callback: Optional[Callable[[str, Dict[str, Any]], None]] = None,
    ) -> Optional[Dict[str, Any]]:
        """Extract a tool definition from (possibly large, multi-page) doc text.

        Splits the text into context-sized chunks, runs the LLM extractor on
        each, and MERGES the results — union of endpoints across chunks, with
        base_url / auth / docs taken from the first chunk that yields them. This
        is what lets a whole crawled doc tree contribute all its endpoints
        instead of only what fits in a single prompt."""
        def _emit(step: str, data: Optional[Dict[str, Any]] = None) -> None:
            if status_callback:
                try:
                    status_callback(step, data or {})
                except Exception:
                    pass

        # Size each chunk + its output to FIT the extraction context window
        # (chars ≈ 3× tokens), leaving room for the system prompt. This keeps a
        # 7B model under 8GB at the default num_ctx (8192) instead of swapping
        # to death on a 32k window — the cap scales up automatically on a
        # machine configured with a larger OLLAMA_EXTRACT_NUM_CTX.
        extract_ctx = settings.OLLAMA_EXTRACT_NUM_CTX
        # Reserve generous output room so a chunk's endpoint JSON is NEVER
        # truncated, then make each chunk deliberately SMALL: fewer endpoints
        # per chunk → the model reliably emits every one and pays full attention
        # to each. More chunks = slower, but the goal here is completeness.
        out_tokens = min(4096, max(2048, extract_ctx // 2))
        # Cap the input slice well under the window; small chunks are the point.
        chunk_tokens = max(800, min(2200, extract_ctx - out_tokens - 1800))
        CHUNK = chunk_tokens * 3
        # Overlap consecutive chunks so an endpoint straddling a boundary is
        # still fully present in at least one chunk (otherwise both halves miss
        # it). ~20% overlap is plenty.
        OVERLAP = CHUNK // 5
        step = max(1, CHUNK - OVERLAP)
        # Process the WHOLE crawled doc — every endpoint must be seen. Each
        # chunk fits the context window, so memory stays bounded no matter how
        # many chunks; only wall-clock time grows with doc size. The upstream
        # crawl's char_cap is the real ceiling.
        chunks = [text[i:i + CHUNK] for i in range(0, len(text), step)] or [text]
        merged_eps: Dict[str, Any] = {}
        base_url = auth_type = docs_url = None
        auth_config: Dict[str, Any] = {}
        rate_limits = examples = None
        merged_quirks: List[str] = []

        for idx, ch in enumerate(chunks):
            _emit("llm_extracting", {"chunk": idx + 1, "of": len(chunks)})
            user_msg = (
                f"Tool name: {tool_name}\nSource: {origin}\n\n"
                f"DOC CONTENT (part {idx + 1} of {len(chunks)}):\n\n{ch}\n\n"
                f"Return the JSON envelope described in the system prompt."
            )
            # Retry once on a transient failure so a single hiccup doesn't
            # silently drop a whole chunk's endpoints.
            part = None
            for attempt in range(2):
                try:
                    part = self._extract_json_smart(
                        _DOCS_EXTRACT_SYSTEM, user_msg,
                        num_predict=out_tokens, num_ctx=extract_ctx,
                        db=db, user_id=user_id,
                        status_callback=status_callback,
                    )
                    break
                except Exception:
                    if attempt == 0:
                        _emit("chunk_retry", {"chunk": idx + 1, "of": len(chunks)})
                        continue
                    logger.warning(
                        "extraction chunk %d/%d failed twice; skipping",
                        idx + 1, len(chunks),
                    )
            if part is None:
                continue
            part = self._normalize_extracted_docs(part or {}, tool_name)
            for k, v in (part.get("endpoints") or {}).items():
                # Prefer the RICHER definition. Bug fix: setdefault used to let
                # an early chunk's bare mention (no body) block a later chunk's
                # full schema — so write-endpoint bodies came back empty.
                merged_eps[k] = (
                    self._merge_endpoint_def(merged_eps[k], v)
                    if k in merged_eps
                    else v
                )
            if not base_url and part.get("base_url") and not _is_placeholder_url(part["base_url"]):
                base_url = part["base_url"]
            if (not auth_type or auth_type == "API_KEY") and part.get("auth_type"):
                auth_type = part["auth_type"]
                if part.get("auth_config"):
                    auth_config = part["auth_config"]
            if not docs_url and part.get("docs_url"):
                docs_url = part["docs_url"]
            if not rate_limits and part.get("rate_limits"):
                rate_limits = part["rate_limits"]
            if not examples and part.get("examples"):
                examples = part["examples"]
            for q in (part.get("quirks") or []):
                if q not in merged_quirks:
                    merged_quirks.append(q)

        if not base_url and not merged_eps:
            return None
        result = {
            "display_name": tool_name.title(),
            "base_url": base_url,
            "auth_type": auth_type,
            "auth_config": auth_config,
            "endpoints": merged_eps,
            "docs_url": docs_url,
            "rate_limits": rate_limits,
            "examples": examples,
            "quirks": merged_quirks or None,
        }
        # Focused second pass so credential labels (Key ID / Key Secret, …) are
        # reliable even when the big extraction skipped them on a small model.
        self._apply_credential_labels(result, tool_name, text, db, user_id)
        return result

    def _upsert_tool(
        self,
        db: Session,
        tool_name: str,
        extracted: Dict[str, Any],
        source_label: str,
    ) -> "ToolDefinition":
        """Insert/update a ToolDefinition from an extracted dict. Idempotent, so
        it's safe to call repeatedly for the incremental per-page crawl saves."""
        fields = dict(
            display_name=extracted.get("display_name") or tool_name.title(),
            base_url=extracted.get("base_url"),
            auth_type=extracted.get("auth_type") or "API_KEY",
            auth_config=extracted.get("auth_config") or {},
            endpoints=extracted.get("endpoints") or {},
            rate_limits=extracted.get("rate_limits"),
            examples=extracted.get("examples"),
            quirks=extracted.get("quirks"),
            docs_url=extracted.get("docs_url"),
            source=source_label,
        )
        row = db.query(ToolDefinition).filter(ToolDefinition.name == tool_name).first()
        try:
            if row:
                for k, v in fields.items():
                    setattr(row, k, v)
                row.last_fetched_at = datetime.utcnow()
            else:
                row = ToolDefinition(name=tool_name, **fields)
                db.add(row)
            db.commit()
        except StaleDataError:
            db.rollback()
            db.query(ToolDefinition).filter(ToolDefinition.name == tool_name).delete()
            db.commit()
            row = ToolDefinition(name=tool_name, **fields)
            db.add(row)
            db.commit()
        db.refresh(row)
        return row

    def _crawl_extract_per_page(
        self,
        db: Session,
        tool_name: str,
        start_url: str,
        *,
        user_id: Optional[int] = None,
        status_callback: Optional[Callable[[str, Dict[str, Any]], None]] = None,
        max_pages: int = 60,
        max_depth: int = 3,
        time_budget_s: float = 600.0,
    ) -> Optional[Dict[str, Any]]:
        """Crawl a doc section and extract each page's endpoint(s) INDIVIDUALLY,
        merging + saving to the DB incrementally — instead of concatenating
        everything into one blob and extracting once.

        Why: doc sites put ~one endpoint per page. Extracting page-by-page gives
        the (small, local) model a focused task per call, so it captures every
        endpoint reliably; the blob approach drops most of them. Saving after
        each page means partial progress persists (and the tool's endpoint count
        grows live) even if the crawl is stopped or times out. Loops over the
        whole discovered section, bounded by max_pages / time."""
        from urllib.parse import urlsplit
        from bs4 import BeautifulSoup
        import time as _time

        def _emit(step: str, data: Optional[Dict[str, Any]] = None) -> None:
            if status_callback:
                try:
                    status_callback(step, data or {})
                except Exception:
                    pass

        base = urlsplit(start_url)
        section = base.path.rstrip("/") or "/"
        deadline = _time.time() + time_budget_s
        visited: set = set()
        queue: List[Tuple[str, int]] = [(start_url, 0)]
        rendered = 0
        acc: Dict[str, Any] = {
            "endpoints": {}, "base_url": None, "auth_type": None,
            "auth_config": {}, "docs_url": None, "rate_limits": None,
            "examples": None, "quirks": [],
        }
        saved_count = -1

        while queue and rendered < max_pages and _time.time() < deadline:
            cur, depth = queue.pop(0)
            if cur in visited:
                continue
            visited.add(cur)
            html = _render_html(cur, settle_ms=1200)
            rendered += 1
            if not html:
                continue
            text = BeautifulSoup(html, "html.parser").get_text("\n")

            if cur == start_url or _has_api_signals(text):
                _emit("extracting_page", {"url": cur, "rendered": rendered})
                # Focused per-page extraction (no status_callback → don't spam
                # per-page chunk counts).
                part = self._extract_from_doc_text(
                    text, tool_name, cur, db=db, user_id=user_id,
                )
                if part:
                    for k, v in (part.get("endpoints") or {}).items():
                        acc["endpoints"][k] = (
                            self._merge_endpoint_def(acc["endpoints"][k], v)
                            if k in acc["endpoints"]
                            else v
                        )
                    if (not acc["base_url"] and part.get("base_url")
                            and not _is_placeholder_url(part["base_url"])):
                        acc["base_url"] = part["base_url"]
                    if (not acc["auth_type"] or acc["auth_type"] == "API_KEY") and part.get("auth_type"):
                        acc["auth_type"] = part["auth_type"]
                        if part.get("auth_config"):
                            acc["auth_config"] = part["auth_config"]
                    if not acc["docs_url"] and part.get("docs_url"):
                        acc["docs_url"] = part["docs_url"]
                    if not acc["rate_limits"] and part.get("rate_limits"):
                        acc["rate_limits"] = part["rate_limits"]
                    if not acc["examples"] and part.get("examples"):
                        acc["examples"] = part["examples"]
                    for q in (part.get("quirks") or []):
                        if q not in acc["quirks"]:
                            acc["quirks"].append(q)

                    # Incremental save once we have something callable, so the
                    # endpoints persist as they're found.
                    if (acc["base_url"] and acc["endpoints"]
                            and len(acc["endpoints"]) != saved_count):
                        self._upsert_tool(
                            db, tool_name,
                            {**acc, "display_name": tool_name.title(),
                             "quirks": acc["quirks"] or None},
                            "user-import",
                        )
                        saved_count = len(acc["endpoints"])
                    _emit("endpoints_so_far", {
                        "count": len(acc["endpoints"]), "pages": rendered,
                    })

            # Harvest more same-section pages (BFS), best-looking links first.
            if depth < max_depth:
                links = self._extract_links_from_html(html, cur, max_links=200)
                cands = {
                    l for l in links
                    if urlsplit(l).netloc == base.netloc
                    and urlsplit(l).path.rstrip("/").startswith(section)
                    and l not in visited
                }
                for l in sorted(cands, key=self._score_api_link, reverse=True):
                    queue.append((l, depth + 1))

        if not acc["endpoints"] and not acc["base_url"]:
            return None
        acc["display_name"] = tool_name.title()
        acc["quirks"] = acc["quirks"] or None
        _emit("crawl_done", {"endpoints": len(acc["endpoints"]), "pages": rendered})
        return acc

    def import_tool_from_source(
        self,
        db: Session,
        tool_name: str,
        *,
        source_url: Optional[str] = None,
        file_bytes: Optional[bytes] = None,
        filename: Optional[str] = None,
        content_type: Optional[str] = None,
        user_id: Optional[int] = None,
        status_callback: Optional[Callable[[str, Dict[str, Any]], None]] = None,
    ) -> Optional[ToolDefinition]:
        """Build a tool definition from a USER-SUPPLIED source instead of web
        discovery. The most reliable path: the user hands us the exact doc.

        Two source shapes (one required):
          • ``source_url``  — a link to an OpenAPI/Swagger spec or a doc page.
          • ``file_bytes``  — an uploaded file (OpenAPI JSON/YAML, or any text/
            markdown/HTML doc). ``filename`` is used only for type hints.

        Strategy, best → fallback:
          1. If the content parses as an OpenAPI/Swagger spec → native parse →
             ALL endpoints, exact (no LLM guessing). This is the 100%-accurate
             path and the whole point of letting users supply the spec.
          2. Otherwise treat it as prose docs → LLM extraction over the text.

        The result is saved as a ToolDefinition with source 'user-import' (or
        'user-import+openapi') so it's never silently overwritten by a web
        refresh of lower quality."""
        tool_name = (tool_name or "").strip().lower()
        if not tool_name:
            return None

        def _emit(step: str, data: Optional[Dict[str, Any]] = None) -> None:
            if status_callback:
                try:
                    status_callback(step, data or {})
                except Exception:
                    logger.exception("status_callback raised; ignoring")

        # ---- 1. Obtain raw text (ANY format: PDF/Word/HTML/JSON/YAML/text) -
        raw_text: str = ""
        # When the source is a JS-heavy HTML doc SITE (not a spec/file), we
        # extract it page-by-page below instead of as one blob.
        per_page = False
        origin_url = source_url or (filename or "uploaded-file")
        if source_url:
            # SSRF guard: we fetch this URL server-side, so block anything that
            # isn't a public http(s) host (internal services, localhost, the
            # cloud metadata endpoint, etc.).
            if not _is_safe_public_url(source_url):
                _emit("error", {
                    "reason": "that link isn't an allowed public http(s) URL "
                    "(internal / localhost / metadata addresses are blocked)"
                })
                logger.warning(f"blocked unsafe source_url for {tool_name}: {source_url!r}")
                return None
            _emit("fetching_source", {"url": source_url})
            try:
                resp = requests.get(
                    source_url, timeout=25,
                    headers={"User-Agent": "Adaptora-DocImport/1.0"},
                )
                resp.raise_for_status()
                # Use raw bytes + content-type so a PDF/Word/HTML URL works too
                # — not just plain-text/JSON URLs.
                ct = resp.headers.get("Content-Type", "")
                raw_text = _extract_text_from_bytes(resp.content, source_url, ct)
                # JS-heavy doc sites (Microsoft Learn, modern API portals) hand
                # back an empty shell over plain HTTP — the real content + nav
                # links are injected by JavaScript. If what we got has no API
                # signals AND isn't already a spec, re-render in headless
                # Chromium and crawl its sub-pages (all server-side, invisible).
                head = raw_text.lstrip()[:200].lower()
                is_specish = head[:1] == "{" or "openapi" in head or "swagger" in head
                if not is_specish and not _has_api_signals(raw_text):
                    # JS-heavy doc site → extract each page individually below
                    # (per-page is far more complete than one big blob).
                    per_page = True
            except Exception as exc:
                _emit("error", {"reason": f"could not fetch source_url: {exc}"})
                logger.warning(f"import fetch failed for {tool_name}: {exc}")
                return None
        elif file_bytes is not None:
            _emit("reading_file", {"filename": filename, "bytes": len(file_bytes)})
            raw_text = _extract_text_from_bytes(
                file_bytes, filename or "", content_type or "",
            )
        if not raw_text.strip():
            _emit("error", {
                "reason": "couldn't read any text from that file/URL "
                "(if it's a scanned image, OCR isn't supported yet)"
            })
            return None

        # ---- 2. Try OpenAPI/Swagger native parse (accurate path) ----------
        spec: Optional[Dict[str, Any]] = None
        stripped = raw_text.lstrip()
        try:
            if stripped.startswith("{"):
                spec = json.loads(raw_text)
            else:
                import yaml  # type: ignore
                loaded = yaml.safe_load(raw_text)
                if isinstance(loaded, dict):
                    spec = loaded
        except Exception:
            spec = None

        extracted: Optional[Dict[str, Any]] = None
        source_label = "user-import"
        if isinstance(spec, dict) and (spec.get("openapi") or spec.get("swagger") or spec.get("paths")):
            _emit("parsing_openapi", {})
            parsed = self._parse_openapi_payload(spec, origin_url)
            if parsed and parsed.get("endpoints"):
                extracted = parsed
                source_label = "user-import+openapi"
                _emit("openapi_parsed", {"endpoints": len(parsed["endpoints"])})

        # ---- 3. Fallback: LLM extraction over the supplied doc -------------
        if extracted is None:
            if per_page and source_url:
                # JS-heavy doc SITE → crawl + extract each page individually,
                # saving incrementally. Captures every endpoint (vs. a blob that
                # drops most) and persists partial progress.
                _emit("crawling_pages", {"url": source_url})
                extracted = self._crawl_extract_per_page(
                    db, tool_name, source_url,
                    user_id=user_id, status_callback=status_callback,
                )
                source_label = "user-import"
            else:
                # A single doc/file (or text that already has API signals) —
                # chunk + merge it.
                extracted = self._extract_from_doc_text(
                    raw_text, tool_name, origin_url,
                    db=db, user_id=user_id, status_callback=status_callback,
                )
            if extracted is None:
                _emit("error", {"reason": "Couldn't extract any endpoints from the doc"})
                return None

        if not extracted or not extracted.get("base_url"):
            _emit("error", {"reason": "could not determine API base_url from the supplied doc"})
            return None
        if not extracted.get("endpoints"):
            _emit("error", {"reason": "no endpoints found in the supplied doc"})
            return None

        # Specs sometimes declare a RELATIVE server url (e.g. Swagger Petstore's
        # "/api/v3"). Resolve it against the source URL's host so calls have a
        # real absolute base. Falls back to the source host root if needed.
        bu = extracted["base_url"]
        if bu.startswith("/") and source_url:
            host = _url_host(source_url)
            if host:
                scheme = "https" if source_url.lower().startswith("https") else "http"
                extracted["base_url"] = f"{scheme}://{host}{bu}".rstrip("/")

        # ---- 4. Save (resilient to concurrent row changes) ----------------
        if extracted.get("docs_url") and _is_placeholder_url(extracted["docs_url"]):
            extracted.pop("docs_url", None)
        display_name = extracted.get("display_name") or tool_name.title()
        row = db.query(ToolDefinition).filter(ToolDefinition.name == tool_name).first()
        fields = dict(
            display_name=display_name,
            base_url=extracted["base_url"],
            auth_type=extracted.get("auth_type") or "API_KEY",
            auth_config=extracted.get("auth_config") or {},
            endpoints=extracted.get("endpoints") or {},
            rate_limits=extracted.get("rate_limits"),
            examples=extracted.get("examples"),
            quirks=extracted.get("quirks"),
            docs_url=extracted.get("docs_url") or (source_url if source_url else None),
            source=source_label,
        )
        try:
            if row:
                for k, v in fields.items():
                    setattr(row, k, v)
                row.last_fetched_at = datetime.utcnow()
            else:
                row = ToolDefinition(name=tool_name, **fields)
                db.add(row)
            db.commit()
        except StaleDataError:
            db.rollback()
            db.query(ToolDefinition).filter(ToolDefinition.name == tool_name).delete()
            db.commit()
            row = ToolDefinition(name=tool_name, **fields)
            db.add(row)
            db.commit()
        db.refresh(row)
        _emit("saved", {"endpoint_count": len(row.endpoints or {}), "auth_type": row.auth_type, "source": source_label})
        return row

    # Targeted query templates run in parallel against the search engine so
    # we cover the four facets the LLM extractor needs: auth/endpoints
    # (official reference), the machine-readable spec (OpenAPI/Swagger),
    # operational limits (rate limits / quotas), and concrete call shapes
    # (code examples). Empirically this surfaces 2-3× more useful chunks
    # than the single-query approach without blowing the prompt budget.
    _DOC_SEARCH_QUERIES: Tuple[str, ...] = (
        "{tool} REST API reference base url authentication endpoints",
        "{tool} OpenAPI swagger specification",
        # GitHub-targeted query — for providers who publish their spec
        # on GitHub (twilio-oai, sendgrid-oai, datadog-api-client, …)
        # the previous queries miss the actual repo URL. This one
        # surfaces it directly.
        "{tool} openapi swagger spec github repository json",
        "{tool} API curl python code example",
        # Extra queries to surface more endpoints from deep-linked doc pages
        "{tool} API endpoints complete list all resources",
        "{tool} developer API documentation full reference",
    )

    def _extract_docs_from_web(
        self,
        tool_name: str,
        *,
        status_callback: Optional[Callable[[str, Dict[str, Any]], None]] = None,
        base_url_hint: Optional[str] = None,
        docs_url_hint: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:
        """Multi-source docs extraction.

        1. Run several targeted web queries in parallel to collect both the
           official reference and supplementary sources (OpenAPI spec,
           rate-limit page, code samples).
        2. If we spot a real OpenAPI/Swagger JSON URL, parse it natively —
           that yields exact endpoints + base URL the LLM can't hallucinate.
        3. Feed the merged page content to the LLM extractor for the rest
           (auth flow, rate limits prose, examples).
        4. Merge native + LLM outputs, with native taking precedence on
           overlap (more reliable)."""
        def _emit(step: str, data: Optional[Dict[str, Any]] = None) -> None:
            if status_callback is not None:
                try:
                    status_callback(step, data or {})
                except Exception:
                    logger.exception("status_callback raised; ignoring")

        _emit("searching_web", {"queries": len(self._DOC_SEARCH_QUERIES)})
        merged_results, engines_used = self._run_multi_source_search(tool_name)
        _emit(
            "web_results",
            {"count": len(merged_results or []), "engines": engines_used},
        )
        if not merged_results:
            logger.info(f"no web results for {tool_name} docs")

        # No URL guessing. The only spec/doc URLs we probe are the ones the
        # web search itself surfaced (real, search-engine-vetted URLs) plus any
        # caller-supplied hints. We never ask the LLM to invent spec URLs.
        llm_base = base_url_hint
        llm_docs = docs_url_hint
        probe_input = list(merged_results or [])

        # Try native OpenAPI parsing on the real URLs found by search — much
        # more reliable than letting the small LLM hallucinate paths from prose.
        openapi_data = self._try_parse_openapi_spec(
            probe_input,
            tool_name,
            base_url_hint=llm_base,
            docs_url_hint=llm_docs,
        )
        if openapi_data and openapi_data.get("endpoints"):
            _emit(
                "openapi_parsed",
                {"endpoints": len(openapi_data["endpoints"])},
            )

        # If we have neither search results nor an OpenAPI spec, there's
        # nothing left for the LLM extractor to chew on — give up cleanly.
        if not merged_results and not openapi_data:
            _emit("error", {"reason": "no usable docs found on the web"})
            return None

        # Phase 1: Enrich initial search results (raw HTML stored for link discovery)
        _emit("enriching", {"max_to_fetch": 10})
        try:
            merged_results = self.llm._enrich_results_with_page_content(
                merged_results,
                max_to_fetch=10,
                per_url_timeout=12.0,
                overall_timeout=45.0,
            )
        except Exception:
            logger.exception("enrichment failed; using raw snippets")

        # Phase 2: Discover additional doc pages via HTML link extraction +
        # sitemap.xml. This handles the "menu navigation" pattern where docs
        # sites spread endpoints across many sub-pages.
        _emit("discovering_pages", {})
        try:
            discovered_urls = self._discover_linked_doc_pages(
                merged_results,
                tool_name,
                docs_url=llm_docs,
                max_pages=30,
            )
        except Exception:
            logger.exception("page discovery failed; skipping")
            discovered_urls = []

        # Deduplicate discovered URLs against what we already have
        existing_urls = {(r.get("href") or "").strip() for r in merged_results}
        new_urls = [u for u in discovered_urls if u not in existing_urls]

        if new_urls:
            _emit("fetching_discovered", {"count": len(new_urls)})
            import concurrent.futures as _cf
            new_results: List[Dict[str, Any]] = []
            pool = _cf.ThreadPoolExecutor(max_workers=min(len(new_urls), 12))
            try:
                future_to_url = {
                    pool.submit(self.llm._fetch_url_content_with_html, url, 10.0): url
                    for url in new_urls[:30]
                }
                done, _ = _cf.wait(future_to_url, timeout=40)
                for fut in done:
                    url = future_to_url[fut]
                    try:
                        result = fut.result()
                    except Exception:
                        result = None
                    if result:
                        text, raw_html = result
                        if text and len(text) > 100:
                            new_results.append({
                                "title": url.split("/")[-1] or url,
                                "href": url,
                                "body": text,
                                "_raw_html": raw_html or "",
                            })
            finally:
                pool.shutdown(wait=False, cancel_futures=True)

            merged_results = merged_results + new_results
            _emit("discovered_fetched", {"new_pages": len(new_results)})

        # Collect doc text for extraction. Kept moderate so a low-RAM host
        # doesn't choke — _extract_from_doc_text chunks this to fit the model's
        # context window anyway. (OpenAPI specs, when found, skip this path.)
        chunks: List[str] = []
        char_budget = 60000
        for r in merged_results[:20]:
            title = (r.get("title") or "").strip()
            url = (r.get("href") or "").strip()
            body = (r.get("body") or "").strip()
            if not (title or body):
                continue
            piece = f"### {title}\nURL: {url}\n\n{body}"
            if char_budget - len(piece) < 0:
                piece = piece[:char_budget]
            chunks.append(piece)
            char_budget -= len(piece)
            if char_budget <= 0:
                break

        if not chunks and not openapi_data:
            _emit("error", {"reason": "no usable page content"})
            return None

        # If a real OpenAPI/Swagger spec was found, it is AUTHORITATIVE: it
        # already gives exact endpoints + request bodies + auth_type. Running the
        # heavy per-chunk LLM extraction on top is pure waste — and on a small
        # local model it's the slow step that makes a refresh look "stuck". So we
        # skip it entirely and let the spec + the focused credential-label pass
        # do the work. The per-chunk LLM extraction only runs when there is NO
        # spec (the model is then the only way to recover endpoints + bodies).
        openapi_has_eps = bool(openapi_data and openapi_data.get("endpoints"))

        extracted: Optional[Dict[str, Any]] = None
        if openapi_has_eps:
            _emit(
                "openapi_authoritative",
                {"endpoints": len(openapi_data["endpoints"])},
            )
        elif chunks:
            _emit(
                "prompt_built",
                {"chunks": len(chunks), "char_budget_used": 100000 - char_budget},
            )
            # Per-chunk extraction + richer-merge (same path the file/URL import
            # uses). A single huge call on a small local model lists endpoints
            # but leaves request BODIES empty; chunking gives each page focused
            # attention and the merge unions body fields across pages — which is
            # what actually populates POST/PUT bodies.
            extracted = self._extract_from_doc_text(
                "\n\n---\n\n".join(chunks),
                tool_name,
                tool_name,
                status_callback=status_callback,
            )

        extracted = self._normalize_extracted_docs(extracted, tool_name)

        # Merge native OpenAPI data over the LLM output. OpenAPI is the
        # ground truth for base_url + endpoints; the LLM owns auth +
        # rate_limits + examples (since those rarely live in the spec).
        if openapi_data:
            if openapi_data.get("base_url"):
                extracted["base_url"] = openapi_data["base_url"]
            if openapi_data.get("endpoints"):
                # OpenAPI endpoints win on key collisions; LLM-only verbs
                # remain available as fallback if the user phrases a request
                # that maps better to the LLM's naming.
                merged_eps = dict(extracted.get("endpoints") or {})
                merged_eps.update(openapi_data["endpoints"])
                extracted["endpoints"] = merged_eps
            if openapi_data.get("docs_url") and not extracted.get("docs_url"):
                extracted["docs_url"] = openapi_data["docs_url"]
            if openapi_data.get("auth_type") and not extracted.get("auth_type"):
                extracted["auth_type"] = openapi_data["auth_type"]

        # We MUST have a base_url to call the API, and it must come from a real
        # fetched doc/spec — either the parsed OpenAPI spec or the base URL the
        # LLM read out of the actual documentation pages. No guessing: if no
        # base_url was found in the docs, there's nothing actionable to save.
        if not extracted.get("base_url"):
            return None
        # Focused second pass for credential labels (the single big web-path
        # extraction often omits them on a small local model).
        self._apply_credential_labels(
            extracted, tool_name, "\n\n".join(chunks)
        )
        return extracted

    # Path segments that should be skipped (marketing / non-reference noise).
    _SKIP_PATH_SIGNALS = re.compile(
        r"/(?:blog|changelog|about|pricing|login|signup|register|status"
        r"|careers|press|legal|terms|privacy|support|community|forum)/",
        re.IGNORECASE,
    )
    # Extensions that are definitely NOT HTML pages.
    _NON_HTML_EXT = re.compile(r"\.(json|yaml|yml|pdf|png|jpg|svg|zip|gz)$", re.IGNORECASE)

    @classmethod
    def _extract_links_from_html(
        cls,
        html: str,
        base_url: str,
        *,
        max_links: int = 60,
    ) -> List[str]:
        """Parse raw HTML and return absolute URLs of pages that look like
        API reference docs. Relative URLs are resolved against base_url.
        Filters out non-HTML, off-domain, and marketing pages."""
        from urllib.parse import urljoin, urlsplit

        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            raw_links = [a.get("href", "") for a in soup.find_all("a", href=True)]
        except Exception:
            # Regex fallback if bs4 fails
            raw_links = re.findall(r'href=["\']([^"\']+)["\']', html)

        base_parts = urlsplit(base_url)
        base_origin = f"{base_parts.scheme}://{base_parts.netloc}"

        out: List[str] = []
        seen: set = set()
        for raw in raw_links:
            raw = raw.strip()
            if not raw or raw.startswith(("#", "mailto:", "javascript:")):
                continue
            # Resolve relative URLs
            absolute = urljoin(base_url, raw)
            parts = urlsplit(absolute)
            # Same domain only
            if parts.netloc != base_parts.netloc:
                continue
            # Strip fragment + query for dedup
            clean = f"{parts.scheme}://{parts.netloc}{parts.path}"
            if cls._NON_HTML_EXT.search(parts.path):
                continue
            if clean in seen or clean == base_url:
                continue
            if cls._SKIP_PATH_SIGNALS.search(parts.path):
                continue
            seen.add(clean)
            out.append(clean)
            if len(out) >= max_links:
                break
        return out

    # Path segments that signal an API-reference page (boosts crawl priority).
    _API_PATH_SIGNALS = re.compile(
        r"/(?:api|rest|reference|endpoint|resource|v\d+|graphql|swagger|openapi"
        r"|methods?|operations?|objects?|types?|schemas?)/",
        re.IGNORECASE,
    )

    @classmethod
    def _score_api_link(cls, url: str) -> int:
        """Heuristic score — higher means more likely to be an API reference
        page. Used to prioritise which linked pages to fetch."""
        score = 0
        if cls._API_PATH_SIGNALS.search(url):
            score += 10
        if re.search(r"/v\d+/", url):
            score += 5
        # Very long paths (deep nested resources) are often individual
        # endpoint pages — worth fetching.
        path_depth = url.count("/")
        score += min(path_depth, 8)
        return score

    def _fetch_sitemap_urls(
        self,
        base_origin: str,
        docs_url: Optional[str],
        tool_name: str,
        *,
        max_urls: int = 40,
    ) -> List[str]:
        """Try fetching sitemap.xml from the docs site and extract API
        reference page URLs. Returns an empty list on any failure."""
        from urllib.parse import urlsplit

        candidates = []
        if docs_url:
            parts = urlsplit(docs_url)
            candidates.append(f"{parts.scheme}://{parts.netloc}/sitemap.xml")
        candidates.append(f"{base_origin}/sitemap.xml")

        for sitemap_url in dict.fromkeys(candidates):
            try:
                resp = requests.get(
                    sitemap_url,
                    timeout=8.0,
                    headers={"User-Agent": "Adaptora/1.0"},
                )
                if resp.status_code != 200:
                    continue
                # Extract all <loc> entries from the sitemap XML
                locs = re.findall(r"<loc>([^<]+)</loc>", resp.text)
                if not locs:
                    continue
                # Filter to API reference pages and score them
                scored = []
                for loc in locs:
                    loc = loc.strip()
                    if self._SKIP_PATH_SIGNALS.search(loc):
                        continue
                    if self._NON_HTML_EXT.search(loc):
                        continue
                    scored.append((self._score_api_link(loc), loc))
                scored.sort(key=lambda x: -x[0])
                result = [url for _, url in scored[:max_urls]]
                if result:
                    logger.info(
                        f"Sitemap {sitemap_url} yielded {len(result)} "
                        f"API reference URLs for {tool_name}"
                    )
                    return result
            except Exception as exc:
                logger.debug(f"sitemap fetch failed for {sitemap_url}: {exc}")
        return []

    def _discover_linked_doc_pages(
        self,
        initial_results: List[Dict[str, Any]],
        tool_name: str,
        *,
        docs_url: Optional[str] = None,
        max_pages: int = 25,
    ) -> List[str]:
        """Discover additional API reference pages by:
          1. Parsing HTML links out of already-fetched pages
          2. Trying sitemap.xml on the docs domain

        Returns a deduplicated, scored list of URLs (highest-relevance first)
        to feed into the content-fetch phase."""
        from urllib.parse import urlsplit

        discovered: Dict[str, int] = {}  # url → score

        # Phase 1: extract links from pages we already fetched
        for r in initial_results:
            href = (r.get("href") or "").strip()
            body_html = (r.get("_raw_html") or "").strip()
            if not href or not body_html:
                continue
            parts = urlsplit(href)
            base_origin = f"{parts.scheme}://{parts.netloc}"
            links = self._extract_links_from_html(body_html, href, max_links=80)
            for link in links:
                if link not in discovered:
                    discovered[link] = self._score_api_link(link)

        # Phase 2: sitemap — highest-signal, covers the whole docs site
        # Use the first search-result domain as the sitemap root.
        if initial_results:
            first_href = (initial_results[0].get("href") or "").strip()
            if first_href:
                parts = urlsplit(first_href)
                base_origin = f"{parts.scheme}://{parts.netloc}"
                sitemap_urls = self._fetch_sitemap_urls(
                    base_origin, docs_url, tool_name, max_urls=max_pages * 2
                )
                for url in sitemap_urls:
                    if url not in discovered:
                        discovered[url] = self._score_api_link(url)

        # Sort by score descending, cap at max_pages
        sorted_urls = sorted(discovered, key=lambda u: -discovered[u])
        return sorted_urls[:max_pages]

    def _run_multi_source_search(
        self, tool_name: str
    ) -> Tuple[List[Dict[str, Any]], List[str]]:
        """Run all _DOC_SEARCH_QUERIES concurrently and return a deduplicated,
        priority-ordered result list plus the set of search engines that
        actually answered. Dedup is by URL — the official reference often
        ranks #1 across multiple queries and we don't want to feed it 4×."""
        import concurrent.futures

        queries = [q.format(tool=tool_name) for q in self._DOC_SEARCH_QUERIES]

        def _run_one(q: str) -> Tuple[List[Dict[str, Any]], str]:
            try:
                results, engine = self.llm._search_web(q, max_results=3)
                return (results or [], engine or "")
            except Exception as exc:
                logger.warning(f"web search '{q}' failed: {exc}")
                return ([], "")

        merged: List[Dict[str, Any]] = []
        seen_urls: set = set()
        engines_used: List[str] = []
        # ThreadPoolExecutor with up to 4 workers — matches len(queries) and
        # plays nice with the underlying HTTP libraries. We use wait() rather
        # than as_completed(timeout=…) because the latter RAISES when any
        # future is still pending at the deadline (killing the good results).
        # Here we want best-effort: take whatever finished in time, drop the
        # slow one(s).
        # Don't use `with ThreadPoolExecutor(...) as pool:` — its __exit__
        # always waits for in-flight futures, so a slow search engine would
        # still block the refresh response by ~10s after our timeout. We
        # shutdown explicitly with wait=False so the response returns as
        # soon as we have enough good results.
        pool = concurrent.futures.ThreadPoolExecutor(max_workers=len(queries))
        try:
            futures = [pool.submit(_run_one, q) for q in queries]
            done, pending = concurrent.futures.wait(
                futures, timeout=25, return_when=concurrent.futures.ALL_COMPLETED
            )
            if pending:
                logger.warning(
                    f"{len(pending)}/{len(futures)} doc-search queries timed "
                    f"out; proceeding with {len(done)} that finished"
                )
            for fut in done:
                try:
                    results, engine = fut.result()
                except Exception as exc:
                    logger.warning(f"search future raised: {exc}")
                    continue
                if engine and engine not in engines_used:
                    engines_used.append(engine)
                for r in results:
                    url = (r.get("href") or "").strip()
                    if not url or url in seen_urls:
                        continue
                    seen_urls.add(url)
                    merged.append(r)
        finally:
            # cancel_futures cancels any not-yet-started; wait=False means
            # we don't block on threads that have already started their HTTP
            # call (those finish in the background and their result is
            # garbage-collected).
            pool.shutdown(wait=False, cancel_futures=True)

        # Rank official/first-party docs to the top and content-farm blogs to
        # the bottom. Downstream steps (OpenAPI probe, page fetch, LLM extract,
        # base_url guess) consume this order, so the official reference gets
        # read first and wins — instead of whichever blog the engine ranked #1.
        # Stable sort: ties keep the engine's original ordering.
        merged.sort(key=lambda r: _score_doc_result(r, tool_name), reverse=True)

        return merged, engines_used

    @staticmethod
    def _looks_like_openapi_url(url: str) -> bool:
        """Heuristic: a URL that's WORTH PROBING as an OpenAPI/Swagger spec.

        Old behaviour required both .json/.yaml suffix AND an
        openapi/swagger keyword in the URL — too strict, missed real
        specs like ``raw.githubusercontent.com/.../api.github.com.json``.

        New rule: any URL that ends in .json/.yaml/.yml. The fetch step
        then validates by looking for ``openapi`` or ``swagger`` keys in
        the JSON; junk JSONs (release notes, blob lists) will fail that
        validation cheaply."""
        if not url:
            return False
        u = url.lower().split("?", 1)[0].split("#", 1)[0]
        return u.endswith(".json") or u.endswith(".yaml") or u.endswith(".yml")

    # Common URL patterns where providers host their OpenAPI/Swagger spec.
    # Filled with a {base} placeholder for the API host. Tried in order;
    # first 200-with-valid-spec wins. Adding more patterns is cheap — each
    # is a single HTTP HEAD-equivalent fetch (8s timeout) and we stop as
    # soon as one matches.
    _OPENAPI_PROBE_PATHS: Tuple[str, ...] = (
        "/openapi.json",
        "/swagger.json",
        "/openapi/v1.json",
        "/api-docs",
        "/v3/api-docs",
        "/.well-known/openapi.json",
        "/swagger/v1/swagger.json",
        "/api/openapi.json",
        "/api/swagger.json",
    )

    @staticmethod
    def _hosts_from_search_results(
        results: List[Dict[str, Any]],
    ) -> List[str]:
        """Pull origins (scheme://host) out of search result URLs. These
        are real, search-engine-vetted hosts — much higher signal than
        templated guesses. Returns unique origins in priority order."""
        from urllib.parse import urlsplit

        out: List[str] = []
        seen: set = set()
        for r in results or []:
            url = (r.get("href") or "").strip()
            if not url:
                continue
            try:
                parts = urlsplit(url if "://" in url else f"https://{url}")
            except Exception:
                continue
            host = (parts.netloc or "").strip()
            if not host:
                continue
            origin = f"{parts.scheme or 'https'}://{host}"
            if origin not in seen:
                seen.add(origin)
                out.append(origin)
        return out

    @classmethod
    def _candidate_probe_urls(
        cls,
        tool_name: str,
        base_url: Optional[str],
        docs_url: Optional[str],
        *,
        search_results: Optional[List[Dict[str, Any]]] = None,
    ) -> List[str]:
        """Build the list of URLs to probe for an OpenAPI spec.

        No guessing — origins come only from real signals:
          1. Explicit hints (a caller-supplied base_url / docs_url).
          2. Hosts extracted from search results (real URLs the search
             engine returned for this tool).

        For each origin, every pattern in _OPENAPI_PROBE_PATHS is joined on.
        Returns a deduplicated list capped to keep latency bounded even when
        no spec exists. If there are no hints and no search hosts, returns []
        (there's nothing real to probe)."""
        from urllib.parse import urlsplit

        hint_origins: List[str] = []
        for u in (base_url, docs_url):
            if not u or not isinstance(u, str):
                continue
            parts = urlsplit(u if "://" in u else f"https://{u}")
            host = parts.netloc.strip()
            if host:
                origin = f"{parts.scheme or 'https'}://{host}"
                if origin not in hint_origins:
                    hint_origins.append(origin)

        search_origins = cls._hosts_from_search_results(search_results or [])

        out: List[str] = []
        seen: set = set()
        # Origins from real hints + search-result hosts, each joined with every
        # probe path. No templated host guesses, no per-tool override URLs —
        # we only probe spec paths on hosts the search actually surfaced.
        for origin in hint_origins + search_origins:
            for path in cls._OPENAPI_PROBE_PATHS:
                url = origin + path
                if url not in seen:
                    seen.add(url)
                    out.append(url)
        # Cap probe count so a misbehaving tool with no spec doesn't pin us
        # against the per-URL timeout × N URLs (8s × N). 24 is enough for
        # 2-3 origins × all probe paths plus one override.
        return out[:24]

    def _fetch_openapi_spec(self, url: str) -> Optional[Dict[str, Any]]:
        """Fetch a single URL and return the parsed JSON if it looks like
        an OpenAPI/Swagger document. None otherwise — both on transport
        failures and on payloads that don't have ``openapi``/``swagger``
        keys. Handles both JSON and YAML; YAML is best-effort (skipped
        silently if PyYAML isn't importable)."""
        url_lower = url.lower().split("?", 1)[0].split("#", 1)[0]
        is_yaml = url_lower.endswith(".yaml") or url_lower.endswith(".yml")
        try:
            resp = requests.get(
                url,
                timeout=8.0,
                headers={
                    "User-Agent": "Adaptora/1.0",
                    "Accept": "application/json,application/yaml,text/yaml,*/*",
                },
            )
            if resp.status_code != 200:
                return None
            ctype = (resp.headers.get("Content-Type") or "").lower()
            looks_json = "json" in ctype or url_lower.endswith(".json")
            looks_yaml = is_yaml or "yaml" in ctype
            if not (looks_json or looks_yaml):
                return None
            # 25 MB cap — GitHub's full spec is ~30 MB but cropped by the
            # parser anyway; bigger payloads are usually not OpenAPI but
            # build artefacts / release-note dumps.
            if int(resp.headers.get("Content-Length") or 0) > 25_000_000:
                return None
            if looks_yaml:
                # PyYAML is best-effort — if it's not installed, skip
                # silently rather than failing the whole refresh.
                try:
                    import yaml  # type: ignore
                except ImportError:
                    logger.debug(
                        f"PyYAML not available; skipping YAML spec at {url}"
                    )
                    return None
                spec = yaml.safe_load(resp.text)
            else:
                spec = resp.json()
        except Exception as exc:
            logger.debug(
                f"openapi fetch failed for {url}: {exc.__class__.__name__}"
            )
            return None
        if not isinstance(spec, dict):
            return None
        # Validate it's actually an OpenAPI / Swagger document by checking
        # for the version key. This cheaply rejects unrelated URLs the
        # search engine surfaced (release notes, blob trees, etc.).
        if not (spec.get("openapi") or spec.get("swagger")):
            return None
        return spec

    # Common filenames/paths providers use when publishing an OpenAPI spec in
    # a GitHub repo — probed by _github_repo_to_spec_urls.
    _GITHUB_SPEC_FILE_PATTERNS: Tuple[str, ...] = (
        "openapi.json",
        "openapi.yaml",
        "openapi.yml",
        "swagger.json",
        "swagger.yaml",
        "spec/openapi.json",
        "spec/openapi.yaml",
        "spec/swagger.json",
        "spec.json",
        "schema/openapi.json",
        "docs/openapi.json",
        "api/openapi.json",
        # twilio-oai-specific shape — multi-API providers who publish per-product
        # specs under a `spec/json/` directory.
        "spec/json/openapi.json",
    )
    _GITHUB_DEFAULT_BRANCHES: Tuple[str, ...] = ("main", "master")

    @classmethod
    def _github_repo_to_spec_urls(
        cls, repo_url: str
    ) -> List[str]:
        """Given a ``https://github.com/{owner}/{repo}`` URL, expand it
        to ``raw.githubusercontent.com`` candidates for every common
        OpenAPI/Swagger spec location. Used to find specs that live in
        GitHub repos (twilio-oai, sendgrid-oai, datadog-api-spec, etc.)
        which the search engine returns as repo links but never as the
        actual raw.githubusercontent.com URL.

        Returns up to ~30 candidates — the probe step caps the actual
        fetches so this can be safely generous."""
        m = re.match(
            r"https?://(?:www\.)?github\.com/([^/\s#?]+)/([^/\s#?]+)",
            repo_url or "",
            re.IGNORECASE,
        )
        if not m:
            return []
        owner, repo = m.group(1), m.group(2)
        # Strip common trailing path segments (.git, /tree/branch).
        repo = re.sub(r"\.git$", "", repo)
        out: List[str] = []
        for branch in cls._GITHUB_DEFAULT_BRANCHES:
            for path in cls._GITHUB_SPEC_FILE_PATTERNS:
                out.append(
                    f"https://raw.githubusercontent.com/{owner}/{repo}/{branch}/{path}"
                )
        return out

    # Regex for any URL that looks like an OpenAPI/Swagger spec file —
    # plain `.json` / `.yaml` URLs we'd want to probe. Used to mine
    # spec links embedded INSIDE search-result bodies (HTML pages often
    # link to `<a href=".../openapi.json">` in the page text).
    _SPEC_URL_RE = re.compile(
        r"https?://[A-Za-z0-9_./\-]+?\.(?:json|yaml|yml)\b",
        re.IGNORECASE,
    )

    @classmethod
    def _mine_spec_urls_from_results(
        cls, results: List[Dict[str, Any]]
    ) -> List[str]:
        """Scan every search-result body for embedded .json/.yaml URLs
        that pass _looks_like_openapi_url. Tavily and Ollama hosted
        search return full page text; trafilatura-enriched results have
        the same. The page content is the BEST signal we have — when
        the spec isn't directly in the URL list, the docs page itself
        usually links to it."""
        out: List[str] = []
        seen: set = set()
        for r in results or []:
            body = r.get("body") or ""
            if not isinstance(body, str):
                continue
            for match in cls._SPEC_URL_RE.findall(body):
                u = match.rstrip(").,;:'\"")
                if u in seen:
                    continue
                seen.add(u)
                out.append(u)
        return out

    @classmethod
    def _github_repo_urls_in_results(
        cls, results: List[Dict[str, Any]]
    ) -> List[str]:
        """Extract unique github.com repo URLs from search-result hrefs
        and bodies. Strips down to the canonical owner/repo form so
        downstream callers can fan out spec file paths without
        duplicates."""
        out: List[str] = []
        seen: set = set()
        repo_re = re.compile(
            r"https?://(?:www\.)?github\.com/([^/\s#?]+)/([^/\s#?]+)",
            re.IGNORECASE,
        )
        for r in results or []:
            for haystack in (r.get("href") or "", r.get("body") or ""):
                if not isinstance(haystack, str):
                    continue
                for m in repo_re.finditer(haystack):
                    canonical = f"https://github.com/{m.group(1)}/{m.group(2)}"
                    canonical = re.sub(r"\.git$", "", canonical)
                    if canonical not in seen:
                        seen.add(canonical)
                        out.append(canonical)
        return out

    def _resolve_external_refs_in_paths(
        self, spec: Dict[str, Any], max_fetches: int = 60
    ) -> Dict[str, Any]:
        """For Swagger 2.0 specs that use EXTERNAL $refs at the path level
        (Mailchimp's spec is the canonical example — every path points
        to a separate JSON file), fetch those refs in parallel and
        inline the resolved content.

        Returns the spec with externally-referenced paths replaced by
        their resolved bodies. Capped at ``max_fetches`` to bound
        latency on extreme cases (Mailchimp has ~180 external refs;
        60 covers the most-used endpoints without blowing past 30s).
        Same-document refs are already handled in ``_parse_openapi_payload``.
        """
        import concurrent.futures

        paths = spec.get("paths")
        if not isinstance(paths, dict) or not paths:
            return spec

        external_refs: List[Tuple[str, str]] = []  # (path_key, ref_url)
        for path_key, ops in paths.items():
            if not isinstance(ops, dict):
                continue
            ref = ops.get("$ref")
            if (
                isinstance(ref, str)
                and ref.lower().startswith(("http://", "https://"))
            ):
                external_refs.append((path_key, ref))
        if not external_refs:
            return spec

        external_refs = external_refs[:max_fetches]
        logger.info(
            f"resolving {len(external_refs)} external $ref(s) in spec…"
        )

        def _fetch(url: str) -> Optional[Dict[str, Any]]:
            try:
                resp = requests.get(
                    url,
                    timeout=6.0,
                    headers={"User-Agent": "Adaptora/1.0"},
                )
                if resp.status_code != 200:
                    return None
                return resp.json()
            except Exception:
                return None

        with concurrent.futures.ThreadPoolExecutor(max_workers=8) as pool:
            future_to_path = {
                pool.submit(_fetch, ref): path_key
                for path_key, ref in external_refs
            }
            try:
                done, _ = concurrent.futures.wait(
                    future_to_path, timeout=30
                )
            except Exception:
                done = []
            for fut in done:
                pk = future_to_path[fut]
                try:
                    resolved = fut.result()
                except Exception:
                    resolved = None
                if isinstance(resolved, dict):
                    paths[pk] = resolved
        return spec

    def _try_parse_openapi_spec(
        self,
        search_results: List[Dict[str, Any]],
        tool_name: str,
        *,
        base_url_hint: Optional[str] = None,
        docs_url_hint: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:
        """Find and parse an OpenAPI/Swagger spec for this tool.

        Tries, in order:
          1. URLs from search results that pattern-match _looks_like_openapi_url
          2. Standard spec paths ({base}/openapi.json, /swagger.json, …) joined
             onto real hosts only — caller-supplied hints + search-result hosts.

        No templated host guesses and no per-tool overrides: every candidate
        comes from a real fetched signal. Returns a partial extracted-docs dict
        (base_url + endpoints + maybe docs_url) on the first valid spec, None if
        none of the candidates produce a usable spec."""
        # 1. Search-result candidates first — these are most likely to be
        # the canonical spec URL (well-indexed, freshly cached).
        search_candidates = [
            (r.get("href") or "").strip()
            for r in (search_results or [])
            if self._looks_like_openapi_url(r.get("href") or "")
        ]
        # 2. Standard spec paths joined onto real hosts (hints + search-result
        #    hosts). No guessing — if there are no real hosts, this is empty.
        probe_candidates = self._candidate_probe_urls(
            tool_name,
            base_url_hint,
            docs_url_hint,
            search_results=search_results,
        )

        seen: set = set()
        ordered: List[str] = []
        for url in search_candidates + probe_candidates:
            if url and url not in seen:
                seen.add(url)
                ordered.append(url)

        if not ordered:
            return None

        # Cap total fetches so latency stays bounded even when no spec exists.
        # 10 attempts × 8s = 80s worst case, but in practice most fail on the
        # cheap status-code / Content-Type check long before a full timeout.
        # Candidates are all real (search-result spec URLs + standard paths on
        # real hosts), so the cap just bounds the standard-path probes.
        for spec_url in ordered[:10]:
            spec = self._fetch_openapi_spec(spec_url)
            if not spec:
                continue
            parsed = self._parse_openapi_payload(spec, spec_url)
            if parsed and parsed.get("endpoints"):
                logger.info(
                    f"Parsed {len(parsed['endpoints'])} endpoints natively "
                    f"from OpenAPI spec at {spec_url} for {tool_name}"
                )
                return parsed
        return None

    @staticmethod
    def _parse_openapi_payload(
        spec: Dict[str, Any], spec_url: str
    ) -> Optional[Dict[str, Any]]:
        """Convert an OpenAPI 3.x (or Swagger 2.x) document into our
        internal endpoints dict. Returns the partial extract or None if the
        spec is unrecognisable. Keeps the parsing minimal — we extract just
        enough for the planner to call the API; the LLM still owns the
        prose-y fields."""
        if not isinstance(spec, dict):
            return None
        paths = spec.get("paths") or {}
        if not isinstance(paths, dict) or not paths:
            return None

        base_url: Optional[str] = None
        # OpenAPI 3.x → servers[].url
        servers = spec.get("servers")
        if isinstance(servers, list) and servers:
            first = servers[0]
            if isinstance(first, dict) and isinstance(first.get("url"), str):
                base_url = first["url"].rstrip("/")
        # Swagger 2.x → host + basePath + schemes
        if not base_url and isinstance(spec.get("host"), str):
            scheme = "https"
            schemes = spec.get("schemes")
            if isinstance(schemes, list) and schemes:
                scheme = schemes[0]
            base_path = spec.get("basePath") or ""
            base_url = f"{scheme}://{spec['host']}{base_path}".rstrip("/")
        if not base_url:
            return None

        # Sort by path length first, then alphabetically. This pushes the
        # most-used surface (e.g. /user, /repos, /issues) into the cap
        # window before deeply-nested admin paths (/enterprises/{enterprise}
        # /…). On GitHub's alphabetical spec, the old order skipped /repos/*
        # entirely because the cap hit during /advisories /agents /apps.
        sorted_paths = sorted(
            paths.items(),
            key=lambda kv: (len(kv[0]), kv[0]),
        )

        # Cap endpoint count. 300 covers a full-featured API surface while
        # keeping DB row size and planner prompts manageable.
        MAX_ENDPOINTS = 300
        endpoints: Dict[str, Dict[str, Any]] = {}
        for path, ops in sorted_paths:
            if not isinstance(ops, dict):
                continue
            # Swagger 2.0 specs (e.g. Mailchimp) often use a single $ref
            # at the path level pointing to another part of the spec
            # where the actual GET/POST/... operations live. Resolve the
            # pointer in-place so the operation loop sees real verbs.
            if "$ref" in ops and isinstance(ops["$ref"], str):
                resolved = _resolve_json_pointer(spec, ops["$ref"])
                if isinstance(resolved, dict):
                    ops = resolved
                else:
                    continue
            for method, op in ops.items():
                if method.lower() not in {"get", "post", "put", "patch", "delete"}:
                    continue
                if not isinstance(op, dict):
                    continue
                # Prefer operationId, fall back to verb-and-tail-path slug.
                op_id = op.get("operationId")
                if not isinstance(op_id, str) or not op_id.strip():
                    op_id = _slug_from_path(method, path)
                op_id = re.sub(r"[^a-zA-Z0-9_]", "_", op_id).strip("_") or _slug_from_path(method, path)

                params_doc: Dict[str, str] = {}
                for p in op.get("parameters") or []:
                    if not isinstance(p, dict):
                        continue
                    if p.get("in") == "query" and isinstance(p.get("name"), str):
                        params_doc[p["name"]] = (p.get("description") or "").strip() or "query parameter"

                description = (op.get("summary") or op.get("description") or "").strip()
                if len(description) > 200:
                    description = description[:200].rsplit(" ", 1)[0] + "…"

                if len(endpoints) >= MAX_ENDPOINTS:
                    break
                endpoints[op_id] = {
                    "method": method.upper(),
                    "path": _normalize_endpoint(path),
                    "description": description or f"{method.upper()} {path}",
                    "params": params_doc or None,
                    # Extract the documented request-body fields from the spec —
                    # this is the ground truth the planner/grounding use to drop
                    # invalid fields ("extra fields sent").
                    "body": _openapi_body_fields(op, spec),
                }
            if len(endpoints) >= MAX_ENDPOINTS:
                break

        out: Dict[str, Any] = {
            "base_url": base_url,
            "endpoints": endpoints,
            "docs_url": spec_url,
        }
        # Best-effort auth_type from securitySchemes.
        comps = spec.get("components") or {}
        sec_schemes = comps.get("securitySchemes") if isinstance(comps, dict) else None
        if not sec_schemes and isinstance(spec.get("securityDefinitions"), dict):
            sec_schemes = spec["securityDefinitions"]
        if isinstance(sec_schemes, dict) and sec_schemes:
            first = next(iter(sec_schemes.values()))
            if isinstance(first, dict):
                t = (first.get("type") or "").lower()
                scheme = (first.get("scheme") or "").lower()
                if t == "oauth2":
                    out["auth_type"] = "OAUTH2"
                elif t == "http" and scheme == "basic":
                    out["auth_type"] = "BASIC"
                elif t == "http" and scheme == "bearer":
                    out["auth_type"] = "BEARER"
                elif t == "apikey":
                    out["auth_type"] = "API_KEY"

        return out

    @staticmethod
    def _merge_endpoints_into(
        target: Dict[str, Dict[str, Any]],
        incoming: Dict[str, Dict[str, Any]],
    ) -> int:
        """Mutate ``target`` to include endpoints from ``incoming``.

        Returns the number of NEW endpoints added. Dedup logic:

        - **By (method, path)**: if ``target`` already has any entry with
          the same (method, path) as an incoming one, skip incoming (it's
          the same endpoint under a possibly different name — keep the
          curated seed name).
        - **By key collision**: if the incoming key is already used by
          ``target`` but points to a DIFFERENT (method, path), keep the
          target entry and add the incoming under a suffixed key
          (``foo``, ``foo_2``, ``foo_3``, …). Lets us preserve both the
          seed's ``list_repos`` and the LLM/spec-discovered ``list_repos``
          when they cover genuinely different endpoints (rare but real).

        Without this, web/spec contributions would either (a) duplicate
        every seed endpoint under a different name, or (b) silently shadow
        each other on key collision. Both result in confusing tool views."""
        # Build the existing (method, path) set once.
        existing_eps: set = set()
        for ep in target.values():
            if not isinstance(ep, dict):
                continue
            m = (ep.get("method") or "").strip().upper()
            p = (ep.get("path") or "").strip()
            if m and p:
                existing_eps.add((m, p))

        added = 0
        for key, ep in (incoming or {}).items():
            if not isinstance(ep, dict):
                continue
            m = (ep.get("method") or "").strip().upper()
            p = (ep.get("path") or "").strip()
            if not p:
                continue
            sig = (m, p) if m else (None, p)
            if sig in existing_eps:
                continue  # same endpoint already present under any key
            # Pick a non-colliding key. Most incoming keys already won't
            # collide; the suffix path only runs for the rare overlap.
            final_key = key
            suffix = 2
            while final_key in target:
                final_key = f"{key}_{suffix}"
                suffix += 1
                if suffix > 20:
                    break  # safety valve
            target[final_key] = ep
            existing_eps.add(sig)
            added += 1
        return added

    # AWS isn't a single API — it's 425+ services. Web search for `aws REST
    # API reference` returns marketing pages, never a useful endpoint list.
    # Instead we introspect boto3's local service model (no network calls
    # required), which knows every service + every operation. We pick a
    # curated list of popular services and their top read-only operations
    # so the row stays manageable (~80 endpoints, not 30,000+).
    #
    # Each entry: service → list of operation_name suffixes. The agent's
    # planner expects endpoints in the form "<service>/<snake_op>", so we
    # build those slugs here.
    _AWS_POPULAR_SERVICES: Tuple[Tuple[str, Tuple[str, ...]], ...] = (
        ("sts",        ("get_caller_identity",)),
        ("ec2",        ("describe_instances", "describe_security_groups",
                        "describe_vpcs", "describe_subnets", "describe_volumes",
                        "describe_images", "describe_key_pairs",
                        "describe_regions", "describe_availability_zones")),
        ("s3",         ("list_buckets", "get_bucket_location",
                        "get_bucket_policy", "list_objects_v2")),
        ("iam",        ("list_users", "list_roles", "list_groups",
                        "list_policies", "get_account_summary")),
        ("lambda",     ("list_functions", "list_layers", "get_account_settings")),
        ("rds",        ("describe_db_instances", "describe_db_clusters",
                        "describe_db_snapshots")),
        ("dynamodb",   ("list_tables", "describe_limits")),
        ("cloudformation", ("list_stacks", "describe_stacks")),
        ("cloudwatch", ("list_metrics", "describe_alarms")),
        ("logs",       ("describe_log_groups", "describe_log_streams")),
        ("sns",        ("list_topics", "list_subscriptions")),
        ("sqs",        ("list_queues",)),
        ("kms",        ("list_keys", "list_aliases")),
        ("secretsmanager", ("list_secrets",)),
        ("ssm",        ("describe_parameters", "describe_instance_information")),
        ("ecs",        ("list_clusters", "list_services", "list_tasks")),
        ("eks",        ("list_clusters", "list_nodegroups")),
        ("route53",    ("list_hosted_zones", "list_health_checks")),
        ("apigateway", ("get_rest_apis",)),
        ("cloudfront", ("list_distributions",)),
    )

    @staticmethod
    def _introspect_aws_endpoints() -> Dict[str, Dict[str, Any]]:
        """Build an additive endpoints dict for the AWS seed by introspecting
        boto3's local service model. Returns {"<verb_slug>": {...}} suitable
        for merging into the existing seed.

        We only emit operations boto3 actually knows about — if a service is
        unavailable in the local boto3 install (older version, etc.) we just
        skip it. No network calls; pure model inspection."""
        try:
            import boto3  # type: ignore
        except ImportError:
            logger.warning("boto3 not installed; skipping AWS introspection")
            return {}

        try:
            session = boto3.Session()
            available = set(session.get_available_services())
        except Exception as exc:
            logger.warning(f"boto3 session/service enumeration failed: {exc}")
            return {}

        endpoints: Dict[str, Dict[str, Any]] = {}
        for service, ops in DynamicAgentService._AWS_POPULAR_SERVICES:
            if service not in available:
                continue
            try:
                # Region is irrelevant for model introspection but required
                # by some service clients. us-east-1 is the safest default.
                client = session.client(service, region_name="us-east-1")
                model_ops = set(client.meta.service_model.operation_names)
            except Exception as exc:
                logger.debug(
                    f"boto3 client init failed for {service}: "
                    f"{exc.__class__.__name__}"
                )
                continue
            for snake_op in ops:
                # boto3 operation_names are PascalCase; convert our snake
                # references to verify the op actually exists.
                pascal_op = "".join(w.capitalize() for w in snake_op.split("_"))
                if pascal_op not in model_ops:
                    continue
                verb_slug = f"{service}_{snake_op}"
                endpoints[verb_slug] = {
                    "method": "POST",  # AWS dispatch goes via boto3, method is moot
                    "path": f"{service}/{snake_op}",
                    "description": f"AWS {service}: {snake_op.replace('_', ' ')}",
                    "params": None,
                    "body": None,
                }
        return endpoints

    @staticmethod
    def _normalize_extracted_docs(
        extracted: Optional[Dict[str, Any]], tool_name: str
    ) -> Dict[str, Any]:
        """Apply the same defensive normalisation the old single-query path
        used: strip hallucinated empties, sanitize URL strings, normalise
        endpoint paths, ensure display_name. Pulled out so the multi-source
        orchestrator and any future caller share the cleanup logic."""
        extracted = _strip_empties(extracted) or {}
        if "endpoints" in extracted and isinstance(extracted["endpoints"], dict):
            extracted["endpoints"] = {
                k: _strip_empties(v)
                for k, v in extracted["endpoints"].items()
                if isinstance(v, dict)
            }
        if isinstance(extracted.get("base_url"), str):
            extracted["base_url"] = _sanitize_url_string(extracted["base_url"])
            # Drop documentation placeholders (api.example.com, your-domain.com,
            # localhost, leftover {templates}). Leaving the key absent makes the
            # caller's `not extracted.get("base_url")` guard fall back to the
            # LLM's real knowledge of the tool instead of saving a dead host.
            if _is_placeholder_url(extracted["base_url"]):
                logger.info(
                    f"discarding placeholder base_url "
                    f"{extracted['base_url']!r} for {tool_name}"
                )
                extracted.pop("base_url", None)
        if isinstance(extracted.get("docs_url"), str):
            extracted["docs_url"] = _sanitize_url_string(extracted["docs_url"])
            if _is_placeholder_url(extracted["docs_url"]):
                extracted.pop("docs_url", None)
        for ep in (extracted.get("endpoints") or {}).values():
            if isinstance(ep, dict) and isinstance(ep.get("path"), str):
                ep["path"] = _normalize_endpoint(ep["path"])
        # rate_limits / examples — keep only if they look structurally valid;
        # the small LLM occasionally emits "rate_limits": "see docs" which is
        # noise. We accept dict for rate_limits and list-of-dicts for examples.
        rl = extracted.get("rate_limits")
        if rl is not None and not isinstance(rl, dict):
            extracted.pop("rate_limits", None)
        ex = extracted.get("examples")
        if ex is not None:
            if not isinstance(ex, list):
                extracted.pop("examples", None)
            else:
                cleaned = [
                    e for e in ex
                    if isinstance(e, dict) and isinstance(e.get("code"), str)
                ]
                extracted["examples"] = cleaned or None
                if not cleaned:
                    extracted.pop("examples", None)
        # quirks — accept a list of non-empty strings only. The small LLM may
        # emit a single string or junk; coerce / drop accordingly.
        qk = extracted.get("quirks")
        if qk is not None:
            if isinstance(qk, str):
                qk = [qk]
            if not isinstance(qk, list):
                extracted.pop("quirks", None)
            else:
                cleaned_q = [
                    s.strip() for s in qk
                    if isinstance(s, str) and s.strip()
                ]
                extracted["quirks"] = cleaned_q or None
                if not cleaned_q:
                    extracted.pop("quirks", None)
        extracted.setdefault("display_name", tool_name.title())
        return extracted

    # =================================================== step 3: connections

    def load_connection(
        self, db: Session, user_id: int, tool_name: str
    ) -> Optional[DynamicToolConnection]:
        return (
            db.query(DynamicToolConnection)
            .filter(
                DynamicToolConnection.user_id == user_id,
                DynamicToolConnection.tool_name == tool_name,
                DynamicToolConnection.is_active == True,  # noqa: E712
            )
            .order_by(DynamicToolConnection.created_at.desc())
            .first()
        )

    def decrypt_credentials(self, conn: DynamicToolConnection) -> Dict[str, Any]:
        if not conn.credentials_encrypted:
            return {}
        try:
            raw = decrypt_api_key(conn.credentials_encrypted)
            return json.loads(raw)
        except Exception:
            logger.exception("failed to decrypt connection credentials")
            return {}

    def save_credentials(
        self,
        db: Session,
        *,
        user_id: int,
        tool: ToolDefinition,
        credentials: Dict[str, Any],
        expires_in_seconds: Optional[int] = None,
    ) -> DynamicToolConnection:
        existing = (
            db.query(DynamicToolConnection)
            .filter(
                DynamicToolConnection.user_id == user_id,
                DynamicToolConnection.tool_name == tool.name,
            )
            .first()
        )
        encrypted = encrypt_api_key(json.dumps(credentials, default=str))
        expires_at = (
            datetime.utcnow() + timedelta(seconds=expires_in_seconds)
            if expires_in_seconds
            else None
        )
        if existing:
            existing.auth_type = tool.auth_type
            existing.display_name = tool.display_name
            existing.credentials_encrypted = encrypted
            existing.token_expires_at = expires_at
            existing.is_active = True
            row = existing
        else:
            row = DynamicToolConnection(
                user_id=user_id,
                tool_name=tool.name,
                display_name=tool.display_name,
                auth_type=tool.auth_type,
                credentials_encrypted=encrypted,
                token_expires_at=expires_at,
            )
            db.add(row)
        db.commit()
        db.refresh(row)
        return row

    def setup_instructions(
        self, tool: ToolDefinition, language: str = "en"
    ) -> Dict[str, Any]:
        """A user-friendly "how to get these credentials" guide.

        Returns ``{"intro": str, "steps": [str, …]}``. For seeded tools the
        instructions are hand-written and live in ``auth_config``. For
        LLM-scraped tools we build a best-effort fallback from the docs
        URL and auth_type so the modal is never empty."""
        cfg = tool.auth_config or {}
        seeded = cfg.get("setup_instructions")
        if isinstance(seeded, dict) and seeded.get("steps"):
            return {
                "intro": seeded.get("intro") or "",
                "steps": list(seeded["steps"]),
            }

        at = (tool.auth_type or "").upper()
        steps: List[str] = []
        intro = (
            f"`{tool.name}` was discovered automatically — these are "
            "generic steps. Check the provider's docs for the exact path."
        )
        if at in ("API_KEY", "BEARER", "PAT"):
            steps = [
                f"Open the provider's developer console: {tool.docs_url or '(see provider docs)'}.",
                "Find the 'API keys' / 'Personal access tokens' / 'Tokens' section and create a new one.",
                "Copy the value the provider shows you (most providers display the secret only ONCE).",
                "Paste it below.",
            ]
        elif at == "BASIC":
            steps = [
                f"Open the provider's developer console: {tool.docs_url or '(see provider docs)'}.",
                "Generate an API key pair — most providers call the two halves 'Key ID' / 'Key Secret' or similar.",
                "Copy both values.",
                "Paste them in the matching fields below.",
            ]
        elif at in ("OAUTH2", "OAUTH2_PKCE", "OAUTH1"):
            steps = [
                f"Register an OAuth application with the provider: {tool.docs_url or '(see provider docs)'}.",
                "Copy the Client ID and Client Secret.",
                "Paste them below and the agent will redirect you to the provider's authorize page next.",
            ]
        elif at == "AWS_SIGV4":
            steps = [
                "Open https://console.aws.amazon.com/iam/home#/users.",
                "Pick an IAM user → Security credentials → Create access key.",
                "Copy Access Key ID + Secret Access Key + the region you want to operate in.",
                "Paste them below.",
            ]
        return {"intro": intro, "steps": steps}

    def required_credential_fields(
        self, tool: ToolDefinition, language: str = "en"
    ) -> List[Dict[str, Any]]:
        """Form schema the frontend renders to collect credentials.

        Different auth types need different fields. For OAUTH2 we ask for
        client_id/secret if not already on file (then the frontend will
        redirect the user through the provider's authorize URL); for
        API_KEY/BEARER/PAT we just need the secret value itself.

        Each tool can override per-field label/placeholder via
        ``auth_config.credential_field_overrides`` — e.g. Razorpay's
        Username/Password become Key ID / Key Secret."""
        at = (tool.auth_type or "API_KEY").upper()
        cfg = tool.auth_config or {}
        overrides = cfg.get("credential_field_overrides") or {}

        def _apply(field: Dict[str, Any]) -> Dict[str, Any]:
            ov = overrides.get(field["name"]) or {}
            return {**field, **ov}

        if at in ("API_KEY", "BEARER", "PAT"):
            label = "API Key" if at == "API_KEY" else ("Bearer Token" if at == "BEARER" else "Personal Access Token")
            placeholder_url = cfg.get("pat_create_url") or tool.docs_url or ""
            placeholder = (
                f"Paste your {tool.display_name} {label.lower()}"
                + (f" (create it here: {placeholder_url})" if placeholder_url else "")
            )
            return [
                _apply(
                    {
                        "name": "secret",
                        "label": label,
                        "type": "password",
                        "required": True,
                        "placeholder": placeholder,
                    }
                )
            ]

        if at == "BASIC":
            # Generic hint: most developer APIs that use HTTP Basic put the
            # API key in the username and the secret in the password. When the
            # docs name them specifically, credential_field_overrides (set at
            # extraction time) replaces these labels with the exact names.
            return [
                _apply({
                    "name": "username", "label": "Username", "type": "text",
                    "required": True,
                    "placeholder": "Usually your API key / Key ID",
                }),
                _apply({
                    "name": "password", "label": "Password", "type": "password",
                    "required": True,
                    "placeholder": "Usually your API secret / Key Secret",
                }),
            ]

        if at == "AWS_SIGV4":
            default_region = (cfg.get("default_region") or "us-east-1")
            return [
                _apply({"name": "access_key_id", "label": "Access Key ID", "type": "text", "required": True}),
                _apply({"name": "secret_access_key", "label": "Secret Access Key", "type": "password", "required": True}),
                _apply({
                    "name": "region",
                    "label": "Region",
                    "type": "text",
                    "required": False,
                    "placeholder": default_region,
                }),
            ]

        if at == "OAUTH1":
            # OAuth 1.0a signs every request with a consumer key/secret +
            # token/secret pair — there's no browser redirect we can drive,
            # so the user pastes all four from the provider's app console.
            return [
                _apply({"name": "consumer_key", "label": "Consumer Key (API Key)", "type": "text", "required": True}),
                _apply({"name": "consumer_secret", "label": "Consumer Secret (API Secret)", "type": "password", "required": True}),
                _apply({"name": "token", "label": "Access Token", "type": "text", "required": True}),
                _apply({"name": "token_secret", "label": "Access Token Secret", "type": "password", "required": True}),
            ]

        if at in ("OAUTH2", "OAUTH2_PKCE"):
            return [
                {
                    "name": "access_token",
                    "label": "Access Token (optional — paste directly to skip OAuth redirect)",
                    "type": "password",
                    "required": False,
                    "placeholder": "Paste a token from the provider's API console to skip the browser flow",
                },
                {
                    "name": "client_id",
                    "label": "OAuth Client ID",
                    "type": "text",
                    "required": False,
                    "placeholder": "Required only if using the Authorize button flow",
                },
                {
                    "name": "client_secret",
                    "label": "OAuth Client Secret",
                    "type": "password",
                    "required": False,
                },
                {
                    "name": "scopes",
                    "label": "Scopes (comma-separated)",
                    "type": "text",
                    "required": False,
                    "placeholder": cfg.get("default_scopes") or "",
                },
            ]

        # Unknown — let the user paste whatever the docs ask for.
        return [
            {
                "name": "secret",
                "label": "Credential",
                "type": "password",
                "required": True,
                "placeholder": f"Paste your {tool.display_name} credential",
            }
        ]

    # =================================================== step 4: plan + run

    @staticmethod
    def _documented_fields(
        endpoints: dict, method: Optional[str], endpoint_path: Optional[str]
    ) -> Tuple[Optional[set], Optional[set]]:
        """Return (param_names, body_names) the DB documents for the endpoint
        matching this method+path, or (None, None) when there's no confident
        match / no documented schema.

        Used to ground the planner: a field that isn't in the documented schema
        is something the model invented, not something the API accepts."""
        norm = _normalize_endpoint(endpoint_path or "")
        if not norm:
            return (None, None)
        want_method = (method or "GET").upper()
        for info in (endpoints or {}).values():
            if not isinstance(info, dict):
                continue
            if (info.get("method") or "GET").upper() != want_method:
                continue
            if _normalize_endpoint(info.get("path") or "") != norm:
                continue
            body = info.get("body")
            params = info.get("params")
            return (
                set(params.keys()) if isinstance(params, dict) and params else None,
                set(body.keys()) if isinstance(body, dict) and body else None,
            )
        return (None, None)

    @staticmethod
    def _filter_endpoints(endpoints: dict, prompt: str, max_endpoints: int = 12) -> dict:
        """Return at most max_endpoints entries most relevant to prompt (keyword match)."""
        if not endpoints or len(endpoints) <= max_endpoints:
            return endpoints
        prompt_lower = prompt.lower()
        keywords = [w for w in prompt_lower.replace("/", " ").split() if len(w) > 2]

        def score(item):
            path, info = item
            text = (path + " " + json.dumps(info, default=str)).lower()
            return sum(1 for kw in keywords if kw in text)

        scored = sorted(endpoints.items(), key=score, reverse=True)
        top = dict(scored[:max_endpoints])
        # always include any exact path match
        for path in endpoints:
            if any(kw in path.lower() for kw in keywords) and path not in top:
                top[path] = endpoints[path]
                if len(top) >= max_endpoints + 4:
                    break
        return top

    def plan_action(
        self,
        *,
        tool: ToolDefinition,
        prompt: str,
        db: Optional[Session] = None,
        user_id: Optional[int] = None,
    ) -> Dict[str, Any]:
        raw_endpoints = tool.endpoints or {}
        filtered = self._filter_endpoints(raw_endpoints, prompt)
        endpoints_summary = json.dumps(filtered, default=str)
        # Provider rules the planner MUST follow (amount units, required
        # headers, date formats, …). These are extracted FROM THIS TOOL'S DOCS
        # at fetch/import time — not hardcoded — so e.g. a payment gateway's
        # "amount is in the smallest currency unit" rule reaches the planner.
        quirks = [q for q in (tool.quirks or []) if isinstance(q, str) and q.strip()]
        quirks_block = (
            "\n\nIMPORTANT PROVIDER RULES (from this API's docs) — follow these "
            "exactly:\n" + "\n".join(f"- {q}" for q in quirks)
            if quirks
            else ""
        )
        user_msg = (
            f"tool: {tool.name}\n"
            f"base_url: {tool.base_url}\n"
            f"endpoints: {endpoints_summary}"
            f"{quirks_block}\n\n"
            f"prompt: {prompt!r}\n\n"
            "Return the JSON envelope described in the system prompt."
        )
        try:
            # Prefer the user's paid model (Claude/GPT) for planning accuracy —
            # complex bodies (payment links, etc.) are exactly where a small
            # local model invents extra/invalid fields. Falls back to the local
            # model when no paid key is configured. A single HTTP-call plan is a
            # small JSON object, so num_predict stays low.
            plan = self._extract_json_smart(
                _ACTION_PLAN_SYSTEM,
                user_msg,
                num_predict=224,
                # Match identify_tool's window so the local model stays warm
                # between identify → plan — a bigger num_ctx forces a reload
                # that stalls the planning step.
                num_ctx=4096,
                db=db,
                user_id=user_id,
            )
        except Exception as exc:
            logger.exception("plan_action: LLM failed")
            return {
                "method": None,
                "endpoint": None,
                "params": None,
                "body": None,
                "summary": f"LLM error: {exc}",
            }

        if isinstance(plan.get("body"), (dict, list)):
            cleaned = _strip_empties(plan["body"])
            plan["body"] = cleaned if cleaned not in (None, {}, []) else None
        if isinstance(plan.get("params"), (dict, list)):
            cleaned = _strip_empties(plan["params"])
            plan["params"] = cleaned if cleaned not in (None, {}, []) else None

        # Reconcile the plan against the DB schema. A small local model often
        # (a) puts body fields into `params` on a POST/PUT/PATCH, and (b)
        # invents fields the API rejects ("extra fields sent"). Using the
        # documented endpoint as ground truth, RELOCATE every provided field to
        # the bucket the docs define it in, and DROP anything documented
        # nowhere. Only runs on a confident endpoint match with a real schema.
        doc_params, doc_body = self._documented_fields(
            raw_endpoints, plan.get("method"), plan.get("endpoint")
        )
        if doc_params or doc_body:
            provided: Dict[str, Any] = {}
            for src in (plan.get("params"), plan.get("body")):
                if isinstance(src, dict):
                    provided.update(src)
            new_params: Dict[str, Any] = {}
            new_body: Dict[str, Any] = {}
            dropped: List[str] = []
            for k, v in provided.items():
                if doc_body and k in doc_body:
                    new_body[k] = v
                elif doc_params and k in doc_params:
                    new_params[k] = v
                else:
                    dropped.append(k)
            if dropped:
                logger.info(
                    "plan_action: dropped undocumented fields for %s %s: %s",
                    tool.name, plan.get("endpoint"), dropped,
                )
            plan["params"] = new_params or None
            plan["body"] = new_body or None
        elif (
            (plan.get("method") or "").upper() in ("POST", "PUT", "PATCH")
            and isinstance(plan.get("params"), dict)
            and plan["params"]
        ):
            # No documented schema to reconcile against, but a write call with
            # everything in the query string is almost always a misplacement —
            # fold params into the JSON body so the provider sees them there.
            plan["body"] = {**(plan.get("body") or {}), **plan["params"]}
            plan["params"] = None

        # Integer-amount providers (Razorpay paise, Stripe cents, PayPal
        # minor units, etc.) reject decimals. Round at the very last step
        # so an LLM that wrote `100.0` instead of `100` doesn't get a
        # cryptic upstream rejection.
        if plan.get("body"):
            plan["body"] = _coerce_integer_money_fields(plan["body"])
        if plan.get("params"):
            plan["params"] = _coerce_integer_money_fields(plan["params"])

        if plan.get("endpoint"):
            plan["endpoint"] = _normalize_endpoint(plan["endpoint"])
        return plan

    def execute_http(
        self,
        *,
        tool: ToolDefinition,
        connection: DynamicToolConnection,
        method: str,
        endpoint: str,
        params: Optional[Dict[str, Any]] = None,
        body: Optional[Any] = None,
        db: Optional[Session] = None,
    ) -> Tuple[int, Any]:
        """Make the authenticated HTTP call. Returns (status_code, parsed_body).

        Branches to ``_execute_boto3`` when the tool's auth_type is
        ``AWS_SIGV4`` — AWS isn't a single REST endpoint so we can't sign
        a generic request; boto3 handles the per-service routing and
        signing for us.

        When ``db`` is supplied and the provider replies 401 on an OAuth2
        connection, we transparently refresh the access token and retry once —
        covers tokens revoked/expired without a known ``expires_at``."""
        if (tool.auth_type or "").upper() == "AWS_SIGV4":
            return self._execute_boto3(
                tool=tool,
                connection=connection,
                endpoint=endpoint,
                kwargs={**(params or {}), **(body or {})} if (params or body) else {},
            )

        # Belt-and-suspenders: even if the DB row pre-dates the extractor
        # sanitization (e.g. a Gmail row with `<https://…>` brackets), strip
        # the junk before requests.request rejects the URL with
        # "No connection adapters were found".
        cleaned_base = _sanitize_url_string(tool.base_url or "")
        cleaned_endpoint = _normalize_endpoint(endpoint)
        if not cleaned_base.lower().startswith(("http://", "https://")):
            raise DynamicAgentError(
                f"Tool `{tool.name}` has an invalid base URL ({tool.base_url!r}). "
                "Re-fetch the docs from the Tools panel or set a seed for this tool."
            )
        url = cleaned_base.rstrip("/") + cleaned_endpoint

        def _send() -> "requests.Response":
            creds = self.decrypt_credentials(connection)
            # url/method/params are only consumed by the OAUTH1 signer.
            headers = self._auth_headers(
                tool, creds, method=method, url=url, params=params
            )
            try:
                return requests.request(
                    method.upper(),
                    url,
                    headers=headers,
                    params=params,
                    json=body if body is not None else None,
                    timeout=30,
                )
            except requests.RequestException as exc:
                raise DynamicAgentError(f"HTTP request to {url} failed: {exc}") from exc

        resp = _send()
        # 401 → token may be stale/revoked. Refresh once and retry.
        if (
            resp.status_code == 401
            and db is not None
            and (tool.auth_type or "").upper() in ("OAUTH2", "OAUTH2_PKCE")
            and self.refresh_oauth_token(db, connection, tool)
        ):
            resp = _send()

        ct = resp.headers.get("content-type", "")
        if "application/json" in ct:
            try:
                parsed: Any = resp.json()
            except ValueError:
                parsed = resp.text
        else:
            parsed = resp.text

        # Some providers (notably Slack) return HTTP 200 with
        # ``{"ok": false, "error": "..."}`` for application-level failures.
        # Translate those into a real non-200 status so the rest of the
        # agent (status badge, summary LLM, logs) treats them as errors.
        effective_status = _interpret_provider_level_error(tool, resp.status_code, parsed)

        connection.last_used_at = datetime.utcnow()
        return effective_status, parsed

    def _execute_boto3(
        self,
        *,
        tool: ToolDefinition,
        connection: DynamicToolConnection,
        endpoint: str,
        kwargs: Dict[str, Any],
    ) -> Tuple[int, Any]:
        """Dispatch an AWS call through boto3.

        ``endpoint`` MUST be ``"<service>/<operation>"`` (e.g.
        ``"ec2/describe_instances"``). ``kwargs`` is forwarded as **kwargs to
        the boto3 method. boto3 raises on credential / parameter errors —
        we translate those into a synthetic HTTP-style ``(status, body)``
        tuple so the rest of the agent doesn't care that this isn't real
        HTTP."""
        try:
            import boto3
            from botocore.exceptions import (
                BotoCoreError,
                ClientError,
                NoCredentialsError,
            )
        except ImportError as exc:
            raise DynamicAgentError(
                "boto3 isn't installed — run `pip install boto3` to enable AWS."
            ) from exc

        creds = self.decrypt_credentials(connection)
        access_key = creds.get("access_key_id") or creds.get("aws_access_key_id")
        secret_key = creds.get("secret_access_key") or creds.get("aws_secret_access_key")
        region = (
            creds.get("region")
            or (tool.auth_config or {}).get("default_region")
            or "us-east-1"
        )
        if not (access_key and secret_key):
            raise DynamicAgentError(
                "AWS access_key_id / secret_access_key are missing from this "
                "connection — re-enter credentials."
            )

        path = (endpoint or "").lstrip("/")
        if "/" not in path:
            raise DynamicAgentError(
                f"AWS endpoint must be '<service>/<operation>', got {endpoint!r}. "
                "Example: 'ec2/describe_instances'."
            )
        service, operation = path.split("/", 1)
        service = service.strip().lower()
        operation = operation.strip()
        if not service or not operation:
            raise DynamicAgentError(
                f"AWS endpoint missing service or operation: {endpoint!r}"
            )

        try:
            client = boto3.client(
                service,
                region_name=region,
                aws_access_key_id=access_key,
                aws_secret_access_key=secret_key,
            )
        except Exception as exc:
            raise DynamicAgentError(
                f"Could not build boto3 client for service `{service}`: {exc}"
            ) from exc

        # ---- Defensive kwarg fixups: small models routinely emit kwargs
        # that boto3 rejects in subtle ways. Normalize before calling so
        # the user doesn't see cryptic AWS errors.
        kwargs = _fixup_boto3_kwargs(service, operation, kwargs, region)

        method_fn = getattr(client, operation, None)
        if not callable(method_fn):
            available = [
                m for m in dir(client) if not m.startswith("_") and callable(getattr(client, m, None))
            ][:30]
            raise DynamicAgentError(
                f"boto3 `{service}` client has no operation `{operation}`. "
                f"A few real operations: {', '.join(available)}…"
            )

        try:
            response = method_fn(**(kwargs or {}))
        except NoCredentialsError as exc:
            connection.last_used_at = datetime.utcnow()
            return 401, {"error": str(exc)}
        except ClientError as exc:
            # AWS's structured error — return as a JSON-shaped body with the
            # real HTTP status so the user sees the actual cause.
            connection.last_used_at = datetime.utcnow()
            err_info = (exc.response or {}).get("Error") or {}
            meta = (exc.response or {}).get("ResponseMetadata") or {}
            return (
                int(meta.get("HTTPStatusCode") or 400),
                {
                    "error": err_info.get("Message") or str(exc),
                    "code": err_info.get("Code"),
                    "service": service,
                    "operation": operation,
                },
            )
        except BotoCoreError as exc:
            raise DynamicAgentError(f"boto3 error: {exc}") from exc

        # boto3 responses sometimes include ResponseMetadata + datetime
        # objects. Strip metadata for the user view and let json default=str
        # handle datetimes upstream.
        if isinstance(response, dict):
            response = {k: v for k, v in response.items() if k != "ResponseMetadata"}
        connection.last_used_at = datetime.utcnow()
        return 200, response

    def _auth_headers(
        self,
        tool: ToolDefinition,
        creds: Dict[str, Any],
        *,
        method: str = "GET",
        url: str = "",
        params: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, str]:
        """Build the auth headers for the upstream call.

        ``method`` / ``url`` / ``params`` are only needed for OAUTH1, whose
        signature is computed per-request over the request line — the other
        auth types ignore them."""
        at = (tool.auth_type or "API_KEY").upper()
        cfg = tool.auth_config or {}
        headers: Dict[str, str] = {
            "Accept": "application/json",
            "User-Agent": "Adaptora-DynamicAgent/1.0",
        }
        # Always merge in any provider-fixed headers (e.g. Notion-Version).
        for k, v in (cfg.get("extra_headers") or {}).items():
            headers[k] = v

        if at in ("API_KEY", "BEARER", "PAT"):
            secret = creds.get("secret") or creds.get("api_key") or creds.get("token")
            if secret:
                header_name = cfg.get("header_name") or "Authorization"
                prefix = cfg.get("credential_prefix")
                if prefix is None:
                    prefix = "Bearer " if at in ("BEARER", "PAT") else ""
                headers[header_name] = f"{prefix}{secret}"
        elif at == "BASIC":
            import base64

            u = creds.get("username") or ""
            p = creds.get("password") or ""
            token = base64.b64encode(f"{u}:{p}".encode()).decode()
            headers["Authorization"] = f"Basic {token}"
        elif at in ("OAUTH2", "OAUTH2_PKCE"):
            token = creds.get("access_token")
            if token:
                header_name = cfg.get("header_name") or "Authorization"
                prefix = cfg.get("credential_prefix") or "Bearer "
                headers[header_name] = f"{prefix}{token}"
        elif at == "OAUTH1":
            ck = creds.get("consumer_key") or creds.get("client_id")
            cs = creds.get("consumer_secret") or creds.get("client_secret")
            tok = creds.get("token") or creds.get("access_token") or ""
            ts = creds.get("token_secret") or ""
            if ck and cs and url:
                headers["Authorization"] = _oauth1_auth_header(
                    method or "GET",
                    url,
                    params,
                    consumer_key=ck,
                    consumer_secret=cs,
                    token=tok,
                    token_secret=ts,
                )
        return headers

    def refresh_oauth_token(
        self, db: Session, conn: DynamicToolConnection, tool: ToolDefinition
    ) -> bool:
        """Use the stored refresh_token to mint a new access_token.

        Returns True on success. Skipped (returns False) when there's no
        refresh_token or the provider's token_url isn't known."""
        creds = self.decrypt_credentials(conn)
        refresh_token = creds.get("refresh_token")
        token_url = (tool.auth_config or {}).get("oauth_token_url")
        client_id = creds.get("client_id")
        client_secret = creds.get("client_secret")
        if not (refresh_token and token_url and client_id and client_secret):
            return False
        try:
            resp = requests.post(
                token_url,
                data={
                    "grant_type": "refresh_token",
                    "refresh_token": refresh_token,
                    "client_id": client_id,
                    "client_secret": client_secret,
                },
                headers={"Accept": "application/json"},
                timeout=20,
            )
            if resp.status_code >= 400:
                logger.warning(
                    f"oauth refresh for {tool.name} returned {resp.status_code}: "
                    f"{resp.text[:200]!r}"
                )
                return False
            body = resp.json()
        except Exception as exc:
            logger.warning(f"oauth refresh for {tool.name} failed: {exc}")
            return False

        new_access = body.get("access_token")
        if not new_access:
            return False
        creds["access_token"] = new_access
        if body.get("refresh_token"):
            creds["refresh_token"] = body["refresh_token"]
        expires_in = body.get("expires_in")
        conn.credentials_encrypted = encrypt_api_key(json.dumps(creds, default=str))
        conn.token_expires_at = (
            datetime.utcnow() + timedelta(seconds=int(expires_in))
            if expires_in
            else None
        )
        db.commit()
        return True

    # =================================================== summary

    def summarize_for_user(
        self,
        *,
        prompt: str,
        plan: Dict[str, Any],
        http_status: Optional[int],
        response_body: Any,
        error: Optional[str],
        language: str = "en",
    ) -> str:
        """Generate the end-user-facing message in the chosen language."""
        sys_prompt = (
            _SUMMARY_HINGLISH_SYSTEM if language == "hinglish" else _SUMMARY_EN_SYSTEM
        )
        if isinstance(response_body, (dict, list)):
            response_excerpt = json.dumps(response_body, default=str)[:1500]
        else:
            response_excerpt = (str(response_body or ""))[:1500]
        user_msg = (
            f"User prompt: {prompt!r}\n"
            f"Action taken: {plan.get('summary') or plan.get('method')} {plan.get('endpoint')}\n"
            f"HTTP status: {http_status}\n"
            f"Error: {error or 'none'}\n"
            f"Response excerpt:\n{response_excerpt}\n\n"
            "Write the user-facing one-paragraph summary now."
        )
        text = _ollama_chat_text(sys_prompt, user_msg, temperature=0.2, num_predict=220)
        # Deterministically replace any hallucinated/placeholder link with the
        # REAL one from the response (extracted from the FULL body, not the
        # truncated excerpt the LLM saw).
        text = _fix_summary_links(text.strip(), response_body)
        return text.strip() or self._fallback_summary(
            language=language, http_status=http_status, error=error
        )

    @staticmethod
    def _fallback_summary(
        *, language: str, http_status: Optional[int], error: Optional[str]
    ) -> str:
        if error:
            if language == "hinglish":
                return f"Action fail ho gaya: {error}"
            return f"Action failed: {error}"
        if http_status is not None and http_status < 400:
            return "Done ✓" if language != "hinglish" else "Ho gaya ✓"
        return f"HTTP {http_status}" if http_status else "Unknown result"

    # =================================================== full turn

    def run_turn(
        self,
        db: Session,
        *,
        user_id: int,
        prompt: str,
        language: str = "en",
        source_url: Optional[str] = None,
        file_bytes: Optional[bytes] = None,
        filename: Optional[str] = None,
        content_type: Optional[str] = None,
        status_callback: Optional[Callable[[str, Dict[str, Any]], None]] = None,
        summarize: bool = True,
    ) -> Dict[str, Any]:
        """End-to-end: identify → docs → connection → plan → execute.

        ``source_url`` (optional) is a user-supplied documentation / OpenAPI
        spec link. ``file_bytes`` (optional) is an uploaded doc/spec file. When
        either is present — or when the prompt itself contains a URL — the
        identified tool's docs are built from THAT source only (no web search),
        mirroring the web UI's import path.

        ``summarize=False`` skips the final user-facing LLM summary call —
        used by the MCP run_action path, where the calling AI assistant
        already turns the raw response into prose, so paying for a third
        local-LLM round-trip just adds latency (and timeout risk). The raw
        http_status / response_body are still returned either way.

        Returns the Thought / Action / Action_Input / Summary envelope the
        frontend renders. On any "blocked" step (no creds, no docs) we
        short-circuit with status=needs_credentials / needs_tool_setup."""
        started = time.perf_counter()
        if language not in ("en", "hinglish"):
            language = "en"

        def emit(step: str, **data: Any) -> None:
            """Best-effort step emit; never let a broken callback break the
            pipeline. Each step keeps the SSE connection alive AND lets the
            UI replace 'Thinking…' with a real label."""
            if not status_callback:
                return
            try:
                status_callback(step, data)
            except Exception as exc:
                logger.warning(f"status_callback({step!r}) failed: {exc}")

        # ---- 1. identify
        emit("identifying_tool", prompt=prompt)
        decision = self.identify_tool(prompt)
        tool_name = decision.get("tool")
        intent = decision.get("intent")
        emit(
            "tool_identified",
            tool=tool_name,
            intent=intent,
            confidence=decision.get("confidence"),
        )

        if not tool_name:
            return self._log_and_return(
                db,
                user_id=user_id,
                language=language,
                prompt=prompt,
                tool_name=None,
                thought=(
                    f"Could not identify a tool from the prompt "
                    f"(reason: {decision.get('reason')})."
                ),
                action="ask_user",
                action_input={"hint": "Tell me which tool to use."},
                status="error",
                summary=(
                    "Mujhe samajh nahi aaya konsa tool use karna hai. "
                    "Tool ka naam ya intent batayein."
                    if language == "hinglish"
                    else "I couldn't tell which tool you meant. "
                    "Try naming the tool (e.g. 'connect github')."
                ),
                started=started,
                final_answer=None,
            )

        # ---- 2. docs
        # Build the tool from a USER-SUPPLIED doc (no web search) when:
        #   • a file was uploaded, or
        #   • an explicit source_url param was given (UI/MCP "import"), or
        #   • the prompt contains a URL AND the router judged this to be a
        #     doc-import request (wants_doc_import). A bare link that is just
        #     content of an action ("post sharing https://…") is NOT treated
        #     as a doc — it falls through to the normal flow.
        explicit_source = (source_url or "").strip() or None
        prompt_url = _first_url_in_text(prompt) if decision.get("wants_doc_import") else None
        doc_source = explicit_source or prompt_url
        if file_bytes:
            emit("importing_doc", tool=tool_name, filename=filename)
            tool = self.import_tool_from_source(
                db, tool_name, file_bytes=file_bytes, filename=filename,
                content_type=content_type, user_id=user_id,
                status_callback=status_callback,
            )
        elif doc_source:
            emit("importing_doc", tool=tool_name, source_url=doc_source)
            tool = self.import_tool_from_source(
                db, tool_name, source_url=doc_source,
                user_id=user_id, status_callback=status_callback,
            )
        else:
            emit("looking_up_docs", tool=tool_name)
            tool = self.lookup_or_fetch_docs(db, tool_name)
        if tool:
            emit(
                "docs_loaded",
                tool=tool.name,
                source=tool.source,
                endpoint_count=len(tool.endpoints or {}),
            )
        if not tool:
            return self._log_and_return(
                db,
                user_id=user_id,
                language=language,
                prompt=prompt,
                tool_name=tool_name,
                thought=(
                    f"Identified `{tool_name}`, but I couldn't find or fetch "
                    "its API docs. Without docs I can't plan a safe call."
                ),
                action="search_docs",
                action_input={"tool": tool_name, "status": "not_found"},
                status="needs_tool_setup",
                summary=(
                    f"`{tool_name}` ka API documentation nahi mil paya. "
                    "Manually setup karne ki zaroorat hai."
                    if language == "hinglish"
                    else f"I couldn't find docs for `{tool_name}`. "
                    "It may need manual setup."
                ),
                started=started,
                final_answer=None,
            )

        thought_id = (
            f"Tool identified: `{tool.name}` ({tool.auth_type}). Docs source: "
            f"`{tool.source}`. Base URL: {tool.base_url}."
        )

        # ---- 3. connection
        emit("checking_connection", tool=tool.name)
        conn = self.load_connection(db, user_id, tool.name)
        if conn:
            emit("connection_found", tool=tool.name, auth_type=tool.auth_type)
        else:
            emit("connection_missing", tool=tool.name, auth_type=tool.auth_type)
        if not conn:
            return self._log_and_return(
                db,
                user_id=user_id,
                language=language,
                prompt=prompt,
                tool_name=tool.name,
                thought=(
                    f"{thought_id} No active credentials for this user — "
                    "must collect them before any call."
                ),
                action="ask_user_creds",
                action_input={
                    "tool": tool.name,
                    "display_name": tool.display_name,
                    "auth_type": tool.auth_type,
                    "credential_fields": self.required_credential_fields(tool, language),
                    "setup_instructions": self.setup_instructions(tool, language),
                    "docs_url": tool.docs_url,
                    "pat_create_url": (tool.auth_config or {}).get("pat_create_url"),
                },
                status="needs_credentials",
                summary=(
                    f"{tool.display_name} se connect karne ke liye credentials "
                    f"chahiye ({tool.auth_type})."
                    if language == "hinglish"
                    else f"To connect {tool.display_name} I need your "
                    f"credentials ({tool.auth_type})."
                ),
                started=started,
                final_answer=None,
            )

        # If intent was clearly "connect" and we already have a connection,
        # just confirm — don't go execute a random action.
        if intent == "connect" or not self._prompt_describes_action(prompt):
            return self._log_and_return(
                db,
                user_id=user_id,
                language=language,
                prompt=prompt,
                tool_name=tool.name,
                thought=(
                    f"{thought_id} Connection already on file — nothing to "
                    "execute, just confirming."
                ),
                action="already_connected",
                action_input={"tool": tool.name},
                status="success",
                summary=(
                    f"{tool.display_name} pehle se connected hai. "
                    "Ab koi action prompt bhejo."
                    if language == "hinglish"
                    else f"{tool.display_name} is already connected. "
                    "Tell me what to do next."
                ),
                started=started,
                final_answer=None,
            )

        # Token-refresh probe for OAuth2 (before the upstream call).
        if (
            (tool.auth_type or "").upper() in ("OAUTH2", "OAUTH2_PKCE")
            and conn.token_expires_at
            and conn.token_expires_at <= datetime.utcnow() + timedelta(seconds=30)
        ):
            emit("refreshing_oauth_token", tool=tool.name)
            self.refresh_oauth_token(db, conn, tool)

        # ---- 4. plan
        emit("planning_action", tool=tool.name)
        plan = self.plan_action(tool=tool, prompt=prompt, db=db, user_id=user_id)
        emit(
            "action_planned",
            method=(plan.get("method") or "").upper(),
            endpoint=plan.get("endpoint"),
            summary=plan.get("summary"),
        )
        method = (plan.get("method") or "").upper()
        endpoint = plan.get("endpoint")
        if not method or not endpoint or method not in {"GET", "POST", "PUT", "PATCH", "DELETE"}:
            return self._log_and_return(
                db,
                user_id=user_id,
                language=language,
                prompt=prompt,
                tool_name=tool.name,
                thought=(
                    f"{thought_id} Planner could not produce a valid action: "
                    f"{plan.get('summary')}"
                ),
                action="plan_action",
                action_input=plan,
                status="error",
                summary=(
                    "Action plan banane mein problem aa gayi. Prompt ko thoda "
                    "specific karke try karo."
                    if language == "hinglish"
                    else "I couldn't build a valid action plan for this prompt. "
                    "Try being more specific."
                ),
                started=started,
                final_answer=None,
            )

        # ---- 5. execute
        emit("executing", tool=tool.name, method=method, endpoint=endpoint)
        try:
            http_status, response_body = self.execute_http(
                tool=tool,
                connection=conn,
                method=method,
                endpoint=endpoint,
                params=plan.get("params"),
                body=plan.get("body"),
                db=db,
            )
            emit("executed", http_status=http_status)
            error_text = (
                None
                if http_status < 400
                else (
                    response_body
                    if isinstance(response_body, str)
                    else json.dumps(response_body, default=str)[:2000]
                )
            )
        except DynamicAgentError as exc:
            http_status = None
            response_body = None
            error_text = str(exc)

        status_label = (
            "success"
            if (http_status is not None and http_status < 400)
            else "error"
        )

        # Once the connection succeeds, mark last_used_at.
        db.commit()

        if summarize:
            emit("summarizing", status=("success" if (http_status is not None and http_status < 400) else "error"))
            summary = self.summarize_for_user(
                prompt=prompt,
                plan=plan,
                http_status=http_status,
                response_body=response_body,
                error=error_text,
                language=language,
            )
        else:
            # Caller (MCP assistant) will summarize the raw response itself —
            # skip the extra local-LLM round-trip. Deterministic fallback
            # keeps a sensible one-liner in the envelope.
            summary = self._fallback_summary(
                language=language, http_status=http_status, error=error_text
            )

        return self._log_and_return(
            db,
            user_id=user_id,
            language=language,
            prompt=prompt,
            tool_name=tool.name,
            thought=(
                f"{thought_id} Plan: {method} {endpoint} — {plan.get('summary')}"
            ),
            action="execute_action",
            action_input={
                "method": method,
                "endpoint": endpoint,
                "params": plan.get("params"),
                "body": plan.get("body"),
            },
            status=status_label,
            summary=summary,
            started=started,
            final_answer=(
                summary
                if status_label == "success"
                else None
            ),
            http_status=http_status,
            response_body=response_body,
            error=error_text,
        )

    def run_endpoint_action(
        self,
        db: Session,
        *,
        user_id: int,
        tool_name: str,
        endpoint_name: str,
        arguments: Optional[Dict[str, Any]] = None,
        language: str = "en",
    ) -> Dict[str, Any]:
        """Fast path for MCP per-endpoint tools.

        The tool AND endpoint are already known (the MCP client invoked a
        specific ``<tool>_<endpoint>`` tool), so we skip the three LLM
        round-trips ``run_turn`` makes — identify_tool, plan_action, and
        summarize_for_user — and execute the HTTP call directly. There is
        no LLM in this path at all, which is what stops the MCP server
        from timing out on a slow local model. The calling AI assistant
        reads the raw response and summarizes it itself."""
        started = time.perf_counter()
        if language not in ("en", "hinglish"):
            language = "en"
        arguments = dict(arguments or {})
        pseudo_prompt = f"{tool_name}.{endpoint_name}"

        # ---- docs (cached → fast)
        tool = self.lookup_or_fetch_docs(db, tool_name)
        if not tool:
            return self._log_and_return(
                db,
                user_id=user_id,
                language=language,
                prompt=pseudo_prompt,
                tool_name=tool_name,
                thought=f"No cached docs for `{tool_name}`.",
                action="search_docs",
                action_input={"tool": tool_name, "status": "not_found"},
                status="needs_tool_setup",
                summary=(
                    f"No docs cached for `{tool_name}`. Run setup_new_tool first."
                ),
                started=started,
                final_answer=None,
            )

        ep = (tool.endpoints or {}).get(endpoint_name)
        if not isinstance(ep, dict):
            return self._log_and_return(
                db,
                user_id=user_id,
                language=language,
                prompt=pseudo_prompt,
                tool_name=tool.name,
                thought=f"Endpoint `{endpoint_name}` not in `{tool.name}` docs.",
                action="execute_action",
                action_input={"tool": tool.name, "endpoint": endpoint_name},
                status="error",
                summary=(
                    f"`{endpoint_name}` isn't a known endpoint for `{tool.name}`. "
                    "Refresh the tool's docs, or use run_action with a "
                    "natural-language prompt."
                ),
                started=started,
                final_answer=None,
            )

        # ---- connection / credentials
        conn = self.load_connection(db, user_id, tool.name)
        if not conn:
            return self._log_and_return(
                db,
                user_id=user_id,
                language=language,
                prompt=pseudo_prompt,
                tool_name=tool.name,
                thought=f"No active credentials for `{tool.name}`.",
                action="ask_user_creds",
                action_input={
                    "tool": tool.name,
                    "display_name": tool.display_name,
                    "auth_type": tool.auth_type,
                    "credential_fields": self.required_credential_fields(tool, language),
                    "setup_instructions": self.setup_instructions(tool, language),
                    "docs_url": tool.docs_url,
                    "pat_create_url": (tool.auth_config or {}).get("pat_create_url"),
                },
                status="needs_credentials",
                summary=(
                    f"To use {tool.display_name} I need your credentials "
                    f"({tool.auth_type})."
                ),
                started=started,
                final_answer=None,
            )

        # OAuth refresh probe (mirror run_turn).
        if (
            (tool.auth_type or "").upper() in ("OAUTH2", "OAUTH2_PKCE")
            and conn.token_expires_at
            and conn.token_expires_at <= datetime.utcnow() + timedelta(seconds=30)
        ):
            self.refresh_oauth_token(db, conn, tool)

        method = (ep.get("method") or "GET").upper()
        endpoint_path = ep.get("path") or ""

        # Substitute {placeholder} path params (e.g. /repos/{owner}/{repo})
        # from the provided arguments; consumed keys won't be re-sent as
        # query/body params.
        used_keys: set = set()

        def _sub(match):
            key = match.group(1)
            if key in arguments and arguments[key] not in (None, ""):
                used_keys.add(key)
                return str(arguments[key])
            return match.group(0)

        endpoint_path = re.sub(r"\{([^}]+)\}", _sub, endpoint_path)

        declared_params = (
            set((ep.get("params") or {}).keys())
            if isinstance(ep.get("params"), dict)
            else set()
        )
        declared_body = (
            set((ep.get("body") or {}).keys())
            if isinstance(ep.get("body"), dict)
            else set()
        )

        params: Dict[str, Any] = {}
        body: Dict[str, Any] = {}
        for k, v in arguments.items():
            if k in used_keys:
                continue
            if k in declared_params:
                params[k] = v
            elif k in declared_body:
                body[k] = v
            elif method in ("GET", "DELETE"):
                params[k] = v
            else:
                body[k] = v

        # ---- execute (no LLM)
        try:
            http_status, response_body = self.execute_http(
                tool=tool,
                connection=conn,
                method=method,
                endpoint=endpoint_path,
                params=params or None,
                body=body or None,
                db=db,
            )
            error_text = (
                None
                if (http_status is not None and http_status < 400)
                else (
                    response_body
                    if isinstance(response_body, str)
                    else json.dumps(response_body, default=str)[:2000]
                )
            )
        except DynamicAgentError as exc:
            http_status = None
            response_body = None
            error_text = str(exc)

        db.commit()
        status_label = (
            "success"
            if (http_status is not None and http_status < 400)
            else "error"
        )
        summary = (
            f"{method} {endpoint_path} → HTTP {http_status}"
            if http_status is not None
            else f"{method} {endpoint_path} failed: {error_text}"
        )

        return self._log_and_return(
            db,
            user_id=user_id,
            language=language,
            prompt=pseudo_prompt,
            tool_name=tool.name,
            thought=(
                f"Direct endpoint call: {method} {endpoint_path} (no LLM planning)."
            ),
            action="execute_action",
            action_input={
                "method": method,
                "endpoint": endpoint_path,
                "params": params or None,
                "body": body or None,
            },
            status=status_label,
            summary=summary,
            started=started,
            final_answer=(summary if status_label == "success" else None),
            http_status=http_status,
            response_body=response_body,
            error=error_text,
        )

    @staticmethod
    def _prompt_describes_action(prompt: str) -> bool:
        """Cheap heuristic: did the user describe an action they want run,
        or did they just say 'connect to <tool>'?"""
        p = (prompt or "").lower()
        connect_only_phrases = ("connect", "setup", "set up", "add ", "link ", "integrate")
        action_verbs = (
            "send", "create", "make", "post", "list", "fetch", "get ", "show",
            "find", "search", "open", "close", "delete", "remove", "update",
            "edit", "comment", "merge", "assign", "schedule", "invite",
            "upload", "download",
        )
        if any(v in p for v in action_verbs):
            return True
        if any(p.startswith(v) or f" {v}" in p for v in connect_only_phrases):
            # Only a connection request, no embedded action verb.
            return False
        # Default: treat as action (user typed a free-form instruction).
        return True

    @staticmethod
    def _log_and_return(
        db: Session,
        *,
        user_id: int,
        language: str,
        prompt: str,
        tool_name: Optional[str],
        thought: str,
        action: str,
        action_input: Any,
        status: str,
        summary: str,
        started: float,
        final_answer: Optional[str],
        http_status: Optional[int] = None,
        response_body: Any = None,
        error: Optional[str] = None,
    ) -> Dict[str, Any]:
        duration_ms = (time.perf_counter() - started) * 1000
        # Truncate response to avoid blowing up the DB row.
        if isinstance(response_body, (dict, list)):
            serialized = json.dumps(response_body, default=str)[:50_000]
        else:
            serialized = (str(response_body) if response_body is not None else None)
            if serialized:
                serialized = serialized[:50_000]

        log = DynamicAgentRunLog(
            user_id=user_id,
            language=language,
            tool_name=tool_name,
            prompt=prompt,
            thought=thought,
            action=action,
            action_input=action_input,
            summary=summary,
            final_answer=final_answer,
            status=status,
            http_status=http_status,
            response_body=serialized,
            error=error,
            duration_ms=duration_ms,
        )
        db.add(log)
        db.commit()
        db.refresh(log)

        return {
            "log_id": log.id,
            "status": status,
            "tool": tool_name,
            "thought": thought,
            "action": action,
            "action_input": action_input,
            "summary": summary,
            "final_answer": final_answer,
            "http_status": http_status,
            "response": response_body,
            "error": error,
            "duration_ms": duration_ms,
            "language": language,
        }

# Module-level singleton — the dependencies (LLMProvider) are cheap to share.
dynamic_agent_service = DynamicAgentService()
